from __future__ import annotations

import argparse
import csv
import json
import math
import re
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt, RGBColor

MAX_WORD_TABLE_COLUMNS = 6
MAX_WORD_TABLE_ROWS = 8
MAX_MAIN_PROGRAM_ASSETS = 6

MOJIBAKE_REPLACEMENTS = {
    "â€™": "'",
    "â€˜": "'",
    "â€œ": '"',
    "â€\x9d": '"',
    "â€\x94": "-",
    "â€”": "-",
    "â€“": "-",
    "â€¦": "...",
    "Ã—": "x",
    "Â ": " ",
    "Â": "",
}

FRIENDLY_COLUMN_HEADERS = {
    "test_key": "Robustness view",
    "scenario_key": "Scenario setting",
    "scenario_label": "Scenario setting",
    "regime": "Regime",
    "versus_series": "Benchmark",
    "risk_adjusted.sharpe_ratio": "Sharpe ratio",
    "return.annualized_return": "Annualized return (%)",
    "static_baseline.excess_return_annualized": "Excess return vs baseline (%)",
    "stress.static_baseline.excess_ann_return": "Stress-period excess return (%)",
    "portfolio.avg_monthly_turnover": "Average turnover (%)",
    "risk.max_drawdown": "Max drawdown (%)",
    "annualized_turnover_ratio": "Annualized turnover (%)",
    "avg_monthly_turnover": "Average monthly turnover (%)",
    "avg_executed_turnover": "Average executed turnover (%)",
    "strategy_ann_return": "Strategy annualized return (%)",
    "versus_ann_return": "Benchmark annualized return (%)",
    "excess_ann_return": "Excess annualized return (%)",
    "strategy_ann_vol": "Strategy volatility (%)",
    "versus_ann_vol": "Benchmark volatility (%)",
    "strategy_max_dd": "Strategy max drawdown (%)",
    "versus_max_dd": "Benchmark max drawdown (%)",
    "hit_rate": "Hit rate (%)",
    "n_periods": "Periods",
    "path_count": "Path count",
    "sharpe_p50": "Median Sharpe",
    "positive_return_probability": "Positive return probability",
    "positive_sharpe_probability": "Positive Sharpe probability",
}

PERCENT_TABLE_HEADERS = {
    "return.annualized_return",
    "static_baseline.excess_return_annualized",
    "stress.static_baseline.excess_ann_return",
    "portfolio.avg_monthly_turnover",
    "risk.max_drawdown",
    "annualized_turnover_ratio",
    "avg_monthly_turnover",
    "avg_executed_turnover",
    "strategy_ann_return",
    "versus_ann_return",
    "excess_ann_return",
    "strategy_ann_vol",
    "versus_ann_vol",
    "strategy_max_dd",
    "versus_max_dd",
    "hit_rate",
}

PROBABILITY_TABLE_HEADERS = {
    "positive_return_probability",
    "positive_sharpe_probability",
    "probability_sharpe_gt_0_50",
    "probability_ann_return_gt_primary_benchmark",
    "oos_hit_rate",
}

INTEGER_TABLE_HEADERS = {"path_count", "n_periods", "scorecard_passed", "scorecard_total"}

SPECIFIC_ASSET_LABELS = {
    "test_01_chart": "Trading-Cost Robustness Chart",
    "test_01_notes": "Trading-Cost Robustness Note",
    "test_01_table": "Trading-Cost Robustness Table",
    "test_02_chart": "Backtest-Window Robustness Chart",
    "test_02_notes": "Backtest-Window Robustness Note",
    "test_02_table": "Backtest-Window Robustness Table",
    "test_04_chart": "Factor-Weight Robustness Chart",
    "test_04_notes": "Factor-Weight Robustness Note",
    "test_04_table": "Factor-Weight Robustness Table",
    "test_05_chart": "Regime-Trigger Robustness Chart",
    "test_05_notes": "Regime-Trigger Robustness Note",
    "test_05_table": "Regime-Trigger Robustness Table",
    "test_06_chart": "Drawdown-Control Robustness Chart",
    "test_06_notes": "Drawdown-Control Robustness Note",
    "test_06_table": "Drawdown-Control Robustness Table",
    "test_07_chart": "Selector-Band Robustness Chart",
    "test_07_notes": "Selector-Band Robustness Note",
    "test_07_table": "Selector-Band Robustness Table",
    "test_08_chart": "Implementation-Constraint Robustness Chart",
    "test_08_notes": "Implementation-Constraint Robustness Note",
    "test_08_table": "Implementation-Constraint Robustness Table",
    "subperiod_regime_decomposition_chart": "Market-Regime Stability Chart",
    "subperiod_regime_decomposition_notes": "Market-Regime Stability Note",
    "subperiod_regime_decomposition_table": "Market-Regime Stability Table",
    "test13_parametric_sharpe_hist": "Simulation-Based Sharpe Distribution Chart",
    "test_13_notes": "Simulation-Based Return Distribution Note",
    "test_13_table": "Simulation-Based Return Distribution Table",
    "acceptance_matrix": "Robustness Coverage Summary",
    "dashboard_notes": "Dashboard Summary Note",
    "stochastic_dashboard": "Simulation-Based Robustness Summary",
    "baseline_scorecard": "Baseline Scorecard Summary",
    "report_summary": "Main-Program Summary",
}


PART_DIR_LABELS = {
    "part_1_deterministic": "Parameter Stability",
    "part_2_ablation": "Component Contribution",
    "part_3_subperiod": "Time Stability",
    "part_4_stochastic": "Simulation-Based Robustness",
    "part_5_dashboard_and_conclusions": "Robustness Summary",
}

PLAIN_METRIC_EXPLANATIONS = [
    "Annualized return describes the strategy's average yearly growth rate over the sample.",
    "Excess return describes how much the strategy beat or lagged its comparison benchmark.",
    "Sharpe ratio describes return earned per unit of overall risk, so a higher value is usually better.",
    "Sortino ratio is similar to Sharpe, but focuses more on harmful downside volatility.",
    "Maximum drawdown describes the worst peak-to-trough loss experienced during the sample.",
    "Turnover describes how much of the portfolio is traded over time, which matters because higher turnover usually means higher implementation cost.",
]

SECTION_DISPLAY_TITLES = {
    "Executive Summary": "Executive Summary",
    "Strategy And Portfolio Construction": "Strategy And Portfolio Construction",
    "Backtest Design": "Backtest Design",
    "Backtest Results": "Backtest Results",
    "Risk, Regime And Exposure Analysis": "Risk, Regime And Exposure Analysis",
    "Robustness And Sensitivity": "Robustness And Sensitivity",
    "Limitations And Monitoring Signals": "Limitations And Monitoring Signals",
    "Product Snapshot": "Product Snapshot",
    "What Changed Since The Last Update": "What Changed Since The Last Update",
    "Performance Update": "Performance Update",
    "Risk And Regime Update": "Risk And Regime Update",
    "Robustness Watch": "Robustness Watch",
    "What Users Should Watch Next": "What Users Should Watch Next",
    "Strategy Summary": "Product Snapshot",
    "Product Design": "How The Investment Process Works",
    "Performance Evaluation": "Performance Update",
    "Robustness Assessment": "Robustness Watch",
    "Risks and Limitations": "Key Risks Users Should Understand",
    "Reporting Priorities": "What Users Should Watch Next",
}

SECTION_INTROS = {
    "Executive Summary": "This summary links the strategy design, latest backtest evidence, main risk caveat, and overall confidence level in one investor-facing view.",
    "Strategy And Portfolio Construction": "This section explains how the portfolio is built and why the process should plausibly create an investable systematic equity exposure.",
    "Backtest Design": "This section sets out the benchmark, cost, execution, rebalance, and metric assumptions that make the backtest interpretable.",
    "Backtest Results": "This section focuses on the realised outcome: absolute return, benchmark-relative value added, downside behaviour, and implementation cost.",
    "Risk, Regime And Exposure Analysis": "This section explains the conditions under which the portfolio did well or struggled, including regime, drawdown, volatility, and exposure concentration.",
    "Robustness And Sensitivity": "This section keeps robustness concise and decision-relevant, using only the checks that materially affect confidence in the main conclusion.",
    "Limitations And Monitoring Signals": "This section states the residual weaknesses and the few signals an investor should watch when reassessing the strategy.",
    "Product Snapshot": "This opening section should help a user understand the product's current state in one pass: what environment it is operating in, whether that environment is supportive or difficult, and the most important point to take away now. It should read like a product-state summary rather than an internal platform-status note.",
    "What Changed Since The Last Update": "This section highlights what is genuinely new or more important now than in the last cycle, so the reader can focus on change rather than rereading a full strategy description.",
    "Performance Update": "This section turns the latest evidence into a few practical conclusions about returns, relative strength, and downside behaviour without forcing the reader through a long list of metrics.",
    "Risk And Regime Update": "This section explains the current market setting, the main risk pressures, and how they may affect user experience of the product right now.",
    "Robustness Watch": "This section explains whether the main message still looks credible once the product is viewed under modest parameter changes, different market slices, and uncertainty rather than under one single historical path. It should explain what the validation means for confidence, not report workflow progress.",
    "What Users Should Watch Next": "This section points to the few signals most worth monitoring from here because they would either strengthen confidence or signal that the current message should be reconsidered. It should sound like guidance for a product user, not a to-do list for the research team.",
    "Strategy Summary": "This opening section should help a user understand the product's current state in one pass: what environment it is operating in, whether that environment is supportive or difficult, and the most important point to take away now. It should read like a product-state summary rather than an internal platform-status note.",
    "Product Design": "This section explains how the investment process works. It is meant to help a non-specialist reader understand where return is expected to come from, how the portfolio is built, and how risk is managed.",
    "Performance Evaluation": "This section turns the latest evidence into a few practical conclusions about returns, relative strength, and downside behaviour without forcing the reader through a long list of metrics.",
    "Robustness Assessment": "This section explains whether the main message still looks credible once the product is viewed under modest parameter changes, different market slices, and uncertainty rather than under one single historical path. It should explain what the validation means for confidence, not report workflow progress.",
    "Risks and Limitations": "This section explains the main conditions under which the product may disappoint, so the reader does not confuse useful historical evidence with a promise of future performance.",
    "Reporting Priorities": "This section points to the few signals most worth monitoring from here because they would either strengthen confidence or signal that the current message should be reconsidered. It should sound like guidance for a product user, not a to-do list for the research team.",
}

SECTION_FALLBACK_ALIASES = {
    "Executive Summary": "Product Snapshot",
    "Backtest Results": "Performance Update",
    "Risk, Regime And Exposure Analysis": "Risk And Regime Update",
    "Robustness And Sensitivity": "Robustness Watch",
    "Limitations And Monitoring Signals": "What Users Should Watch Next",
}

SECTION_PARAGRAPH_EXCLUSIONS = {
    "Product Snapshot": [
        "data health",
        "pipeline",
        "dag",
        "artifact",
        "artifacts",
        "delivery includes",
        "row-level persistence",
        "freshness sla",
        "csv datasets",
        "manifest",
        "workflow",
        "all key checks passing",
        "coverage is comprehensive",
        "ongoing data health",
        "handoff pack",
        "timestamped",
        "artifacts delivered",
        "delivery artifacts",
        "robustness handoff",
        "checks passing",
        "health status",
    ],
    "Robustness Watch": [
        "acceptance",
        "partial acceptance",
        "remaining block",
        "pass count",
        "robustness blocks passed",
        "acceptance matrix",
        "partial robustness block",
        "deterministic tests",
        "stochastic tests",
        "ablation blocks",
        "completed successfully",
        "completed.",
        "completed,",
        "completed",
        "all deterministic",
        "all ablation",
        "all stochastic",
        "partial implementation",
        "missing rows",
        "partially done",
    ],
    "Risk And Regime Update": [
        "quality (",
        "value (",
        "market technical (",
        "dividend (",
        "dividend tilt",
        "momentum shift",
        "factor exposures remain stable",
        "factor exposures",
        "unchanged factor exposures",
    ],
    "What Users Should Watch Next": [
        "future updates should",
        "follow-up",
        "needs resolution",
        "should be completed",
        "research team",
        "internal next step",
        "future work",
        "team task",
        "data health",
        "dag health",
        "partial robustness block",
        "ongoing monitoring is needed",
        "should focus on",
        "requires emphasis",
        "research priority",
        "next step",
        "deterministic tests",
        "stochastic tests",
        "partial acceptance",
        "partial block",
        "acceptance block",
        "fixed window subperiod tables",
        "recommend completing",
        "outstanding",
    ],
}

PART_SUMMARIES = {
    "Part 1": "Parameter-stability evidence on a quarterly-rebalanced strategy with monthly performance measurement, so the report can show whether the mainline result survives practical changes.",
    "Part 2": "Ablation evidence that isolates specific building blocks and shows what each block contributes to the final result rather than only reporting the combined strategy.",
    "Part 3": "Subperiod evidence that checks whether behaviour is stable across fixed windows and regime-based slices instead of only over the full sample.",
    "Part 4": "Simulation-based robustness evidence using bootstrap, perturbation, neighbourhood, and path-style checks so uncertainty is discussed explicitly.",
    "Part 5": "Final dashboard and conclusions evidence that pulls the robustness outputs together into report-ready takeaways.",
}

PART_REPORT_USE = {
    "Part 1": "This evidence shows whether the headline result still stands after practical implementation settings are nudged away from the baseline.",
    "Part 2": "This evidence shows which building blocks are genuinely carrying the result, rather than letting the full process hide weak components.",
    "Part 3": "This evidence shows whether the result is broad-based through time rather than confined to one especially favourable historical window.",
    "Part 4": "This evidence shows whether the main conclusion still looks believable when randomness and alternative paths are introduced.",
    "Part 5": "This evidence turns the broader robustness work into a small set of decision-relevant takeaways for a product user.",
}

ITEM_SUMMARIES = {
    "Test 1": "Checks whether the strategy still looks attractive when realistic trading-cost assumptions are pushed higher or lower.",
    "Test 2": "Checks whether the headline result changes materially when the backtest starts earlier or later.",
    "Test 3": "Checks whether the strategy depends too heavily on how concentrated the final portfolio is.",
    "Test 4": "Checks whether the result still holds after small changes to the factor-weight mix rather than one exact recipe.",
    "Test 5": "Checks whether the regime-trigger setting has to be tuned very precisely for the strategy to work.",
    "Test 6": "Checks the trade-off between stronger downside protection and the cost of applying that protection.",
    "Test 7": "Checks whether the selector band is helping to stabilise the portfolio rather than adding noise.",
    "Test 8": "Checks whether tighter real-world implementation constraints materially weaken the investment case.",
    "Ablation A": "Checks what happens when one major building block is removed from the full strategy.",
    "Ablation B": "Checks whether the second major building block contributes real value to the final result.",
    "Ablation C": "Checks whether the third major building block is necessary for the strategy to remain convincing.",
    "Fixed windows": "Checks whether the result still looks credible across different historical windows rather than one favourable sample.",
    "Regime decomposition": "Checks where the strategy performs well, where it struggles, and how this changes between normal and stressed markets.",
    "test_9": "Checks whether the result still looks credible when the realised return path is resampled many times.",
    "test_10": "Checks whether the strategy remains worthwhile under less favourable trading-cost assumptions.",
    "test_11": "Checks whether nearby factor-weight choices still lead to the same broad conclusion.",
    "test_12": "Checks whether the strategy continues to work outside the original in-sample fit.",
    "test_13": "Checks whether the investment case still holds under simulated return paths rather than only the realised history.",
    "Coverage note": "Explains the intended interpretation and caveats for the surrounding evidence block.",
    "Stochastic report ready": "Collects higher-level stochastic notes and summaries for direct use in the written report.",
}

ITEM_REPORT_USE = {
    "Test 1": "Use this test to explain whether the strategy depends on unrealistically low trading friction.",
    "Test 2": "Use this test to explain whether the result depends on one lucky sample start date.",
    "Test 3": "Use this test to explain whether the strategy only works at one particular concentration setting.",
    "Test 4": "Use this test to explain whether the factor design is robust or over-tuned.",
    "Test 5": "Use this test to explain whether the regime logic is robust or fragile.",
    "Test 6": "Use this test to explain whether downside protection helps enough to justify its cost.",
    "Test 7": "Use this test to explain whether the selector band is adding genuine stability.",
    "Test 8": "Use this test to explain whether the strategy still survives practical portfolio-construction limits.",
    "Ablation A": "Use this ablation to explain what the first major building block contributes to the final strategy.",
    "Ablation B": "Use this ablation to explain what the second major building block contributes to the final strategy.",
    "Ablation C": "Use this ablation to explain what the third major building block contributes to the final strategy.",
    "Fixed windows": "Use this evidence to explain whether the result looks stable through time.",
    "Regime decomposition": "Use this evidence to explain in which market conditions the strategy earns its edge and in which conditions it struggles.",
    "test_9": "Use this test to explain whether the result still looks credible after resampling the path.",
    "test_10": "Use this test to explain whether the strategy remains investable when costs worsen.",
    "test_11": "Use this test to explain whether nearby factor settings still support the same conclusion.",
    "test_12": "Use this test to explain whether the result generalises outside the original fit period.",
    "test_13": "Use this test to explain whether the investment case remains credible under simulated paths.",
    "Coverage note": "Use this note to explain scope, caveats, and what the surrounding evidence is intended to prove.",
    "Stochastic report ready": "Use this block to pull the stochastic evidence into polished, report-ready interpretation rather than raw test output only.",
}

PART_SELECTION_LIMITS = {
    "Part 1": 1,
    "Part 2": 0,
    "Part 3": 1,
    "Part 4": 1,
    "Part 5": 0,
}

MAIN_REPORT_SUMMARIES = {
    "nav_vs_benchmarks": "Compares the strategy against the benchmark and baseline so the report can explain absolute and relative performance in the same place.",
    "drawdown_comparison": "Highlights the depth and timing of drawdowns so the report can discuss resilience alongside return.",
    "turnover_and_cost": "Links implementation activity to cost drag and feasibility.",
    "regime_return_summary": "Shows how the strategy behaves across normal and stress regimes, which is useful when explaining the regime-aware design.",
    "latest_sector_risk_contribution": "Shows where current risk concentration sits at sector level.",
    "trade_blotter": "Provides a compact preview of implementation activity and the latest trade list.",
    "report": "Captures the latest markdown report bundle for direct textual reference.",
    "report_summary": "Provides a concise supporting summary of the latest main-program reporting run.",
}

MAIN_REPORT_USE = {
    "nav_vs_benchmarks": "Use this figure to explain headline product quality in one place: overall growth, benchmark-relative performance, and whether the active process improved on a simpler baseline.",
    "drawdown_comparison": "Use this figure to discuss downside behaviour and whether the product delivered a more tolerable loss profile than its reference portfolios.",
    "turnover_and_cost": "Use this figure to connect portfolio activity to feasibility. It helps explain whether observed returns rely on unrealistic trading intensity.",
    "regime_return_summary": "Use this figure to explain conditional performance and why the regime-aware design matters to the user experience of the product.",
    "latest_sector_risk_contribution": "Use this figure to explain where current risk concentration sits and whether the product remains diversified enough for its stated role.",
    "trade_blotter": "Use this table to illustrate how the current portfolio is being implemented in practice, including what is changing and why.",
    "report": "Use this supporting text only as background material for consistency checks against the generated narrative.",
    "report_summary": "This summary keeps a compact supporting snapshot of the latest reporting run without repeating the main narrative.",
}

MAIN_REPORT_INVESTOR_INTERPRETATIONS = {
    "report_summary": (
        "For an investor, the key read-through is that the strategy has a positive active-return story after the stated implementation assumption, "
        "but it is still an equity strategy with meaningful drawdown and trading intensity. That makes the conclusion investable but conditional: "
        "confidence depends on the return edge remaining large enough to absorb future cost and regime pressure."
    ),
    "nav_vs_benchmarks": (
        "The NAV path is the clearest evidence for the strategy's portfolio role. It supports presenting the product as a growth-oriented active equity sleeve "
        "rather than a defensive substitute, because the benefit is mainly benchmark-relative compounding rather than loss avoidance."
    ),
    "drawdown_comparison": (
        "Drawdown evidence matters because it describes the holding experience an investor would actually face. The chart supports a controlled-equity-risk message "
        "only if drawdowns remain comparable with the stated references, not if the report promises protection."
    ),
    "turnover_and_cost": (
        "This is the main implementation checkpoint. The strategy can be attractive if excess return continues to survive costs, but capacity, liquidity, and slippage "
        "deserve more attention than the headline NAV chart alone would suggest."
    ),
    "regime_return_summary": (
        "The regime evidence keeps the investment case honest. It can support relative resilience in stress, but it should be framed as conditional performance rather "
        "than as a promise that the product will protect capital in every difficult market."
    ),
    "latest_sector_risk_contribution": (
        "This evidence is most useful for risk governance. It helps investors see whether current risk is concentrated in one sector, while leaving factor-quality "
        "claims to the dedicated factor and attribution evidence."
    ),
}

ROBUSTNESS_INVESTOR_INTERPRETATIONS = {
    "Part 1": (
        "For investors, parameter tests are confidence qualifiers rather than a second investment thesis. Stable results across reasonable settings make the backtest "
        "less fragile; sharp deterioration would turn the result into a tuning concern."
    ),
    "Part 3": (
        "This evidence addresses sample dependence. If the conclusion survives different windows or regimes, the strategy reads less like one favourable historical "
        "episode and more like a repeatable process; weak segments should still be named as monitoring risks."
    ),
    "Part 4": (
        "The uncertainty checks help translate one realised backtest into a range of plausible outcomes. That is useful for setting confidence and expectations, "
        "but it should not be mistaken for a guarantee that future paths will resemble the central case."
    ),
}

MAIN_REPORT_INVESTOR_OVERRIDE_KEYS = {
    "report_summary",
    "turnover_and_cost",
    "regime_return_summary",
    "latest_sector_risk_contribution",
}

IMAGE_NOTE_KEYWORDS = {
    "nav": "This chart explains the trajectory of the strategy or reference series over time.",
    "drawdown": "This chart explains downside episodes and where the deepest peak-to-trough losses occurred.",
    "turnover": "This chart explains how implementation activity and cost interact with realised performance.",
    "cost": "This visual explains execution drag and the cost burden attached to the strategy.",
    "regime": "This chart explains how behaviour changes between normal and stress conditions.",
    "risk": "This visual explains where current concentration or contribution to risk is coming from.",
    "sharpe": "This chart explains risk-adjusted performance rather than raw return alone.",
    "heatmap": "This visual explains period-by-period variation and consistency.",
    "correlation": "This chart explains co-movement, diversification, or clustering.",
    "bootstrap": "This visual explains resampling robustness and uncertainty around realised results.",
    "monte": "This visual explains simulated outcome distributions under alternative paths.",
    "oos": "This visual explains out-of-sample behaviour and generalisation.",
    "sample_paths": "This visual explains how simulated paths can diverge from the realised history.",
}

BRIEF_IMAGE_CAPTIONS = {
    "nav": "Portfolio path relative to the benchmark and baseline.",
    "drawdown": "Peak-to-trough drawdown profile over the sample.",
    "turnover": "Turnover and implementation-cost profile.",
    "cost": "Performance sensitivity to trading-cost assumptions.",
    "regime": "Comparative behaviour across normal and stress conditions.",
    "risk": "Current distribution of portfolio risk contribution.",
    "sharpe": "Risk-adjusted performance over time.",
    "heatmap": "Month-by-month variation in realised returns.",
    "correlation": "Correlation structure and diversification relationships.",
    "bootstrap": "Resampled outcome distribution under bootstrap perturbations.",
    "monte": "Distribution of outcomes under simulated return paths.",
    "oos": "Rolling out-of-sample performance through time.",
    "sample_paths": "Illustrative simulated return paths from the stochastic exercise.",
}

SPECIFIC_IMAGE_CAPTIONS = {
    "test_01_chart": (
        "The x-axis shows progressively higher assumed trading-cost levels, from 10 bps to 40 bps, "
        "with 15 bps marking the mainline implementation assumption."
    ),
    "test_02_chart": (
        "The x-axis shifts the backtest start point around the mainline window so the reader can see "
        "whether the result depends too heavily on one specific sample start."
    ),
    "test_04_chart": (
        "Each x-axis label denotes one factor-weight perturbation scenario. "
        "For example, 'Dividend -5pp' means the dividend factor weight is reduced by five percentage points, "
        "'Quality +5pp' means that factor weight is increased by five percentage points, "
        "and 'Equal Weight' marks the neutral comparison case."
    ),
    "test_05_chart": (
        "The x-axis moves the regime trigger from disabled through less sensitive, base, and more sensitive settings, "
        "so the chart shows how much the switching rule affects Sharpe and drawdown."
    ),
    "test_06_chart": (
        "The x-axis orders the drawdown brake from off to increasingly forceful variants, showing the trade-off "
        "between protection and implementation drag."
    ),
    "test_07_chart": (
        "The x-axis tightens the selector band from none through wide and medium to narrow, showing how strongly "
        "selection smoothing changes performance and turnover."
    ),
    "test_08_chart": (
        "The x-axis tightens trading constraints from none to strong. The near-flat profile indicates that the mainline "
        "result is not highly dependent on this setting within the tested range."
    ),
}

SPECIFIC_TABLE_CAPTIONS = {
    "test_01_table": "This table shows how the headline performance measures change as trading-cost assumptions rise.",
    "test_02_table": "This table shows whether the result changes materially when the backtest starts earlier or later.",
    "test_04_table": "This table compares the strategy after each factor-weight adjustment so the reader can see which weights matter most.",
    "test_05_table": "This table compares performance under different regime-trigger settings to show how sensitive the switching rule is.",
    "test_06_table": "This table compares the payoff from weaker or stronger drawdown-protection settings.",
    "test_07_table": "This table shows how the selector band changes performance and implementation behaviour.",
    "test_08_table": "This table shows whether tighter implementation constraints materially change the investment case.",
}

PLAIN_LANGUAGE_EVIDENCE_NOTES = {
    "NAV Vs Benchmarks": "This chart shows whether the product has continued to outperform its benchmark and baseline over the full period.",
    "Drawdown Comparison": "This chart shows how severe the main loss episodes were and whether the product recovered in a reasonable way.",
    "Turnover And Cost": "This chart shows whether the product still looks practical once realistic trading friction is taken into account.",
    "Regime Return Summary": "This chart shows whether the product behaves differently in calmer markets and in stressed markets.",
    "Latest Sector Risk Contribution": "This chart shows where the portfolio's current risk is concentrated.",
    "Test 1": "This check shows whether the product still looks attractive when trading costs rise.",
    "Test 2": "This check shows whether the result still holds if the sample starts earlier or later.",
    "Test 3": "This check shows whether the product still works when the portfolio is made broader or more concentrated.",
    "Test 4": "This check shows whether the result depends on one very specific factor mix or still holds after small weight changes.",
    "Test 5": "This check shows whether the regime trigger has to be tuned very precisely for the strategy to work.",
    "Test 6": "This check shows how much downside protection helps and what it costs in return or turnover.",
    "Test 7": "This check shows whether the selector band is genuinely stabilising the portfolio or simply adding noise.",
    "Test 8": "This check shows whether the investment case survives tighter real-world implementation constraints.",
    "Ablation A": "This ablation shows what happens if one important building block is removed from the full strategy.",
    "Ablation B": "This ablation shows whether the second building block contributes meaningful value to the final result.",
    "Ablation C": "This ablation shows whether the third building block is necessary for the strategy to remain convincing.",
    "Fixed Windows": "This section shows whether the strategy still looks credible across different historical windows rather than only one favourable span.",
    "Regime Decomposition": "This section shows where the strategy does well, where it struggles, and how that changes between normal and stressed conditions.",
    "test_9": "This simulation-based check shows whether the result still looks credible when the realised path is resampled many times.",
    "test_10": "This simulation-based check shows whether the strategy remains worthwhile when trading costs become less favourable.",
    "test_11": "This simulation-based check shows whether nearby factor-weight choices still lead to the same broad conclusion.",
    "test_12": "This rolling check shows whether the strategy keeps working outside the original in-sample fit.",
    "test_13": "This simulation-based check shows whether the investment case still holds under simulated return paths rather than just the realised history.",
}

PREFERRED_MAIN_REPORT_ORDER = [
    "report_summary.json",
    "nav_vs_benchmarks.png",
    "drawdown_comparison.png",
    "turnover_and_cost.png",
    "regime_return_summary.png",
    "latest_sector_risk_contribution.png",
]

APPENDIX_MAIN_REPORT_ORDER = []

MAIN_REPORT_CORE_EXCLUSIONS = set()


INTERNAL_NOTE_PREFIXES = (
    "interpretation:",
    "source table:",
    "figure:",
    "extra reference nav chart:",
    "source file:",
    "writing notes",
    "use this ",
    "use these ",
    "use the ",
    "the report should ",
    "in the written report",
    "this table is the cleanest source",
    "selected columns are shown",
    "table section ",
    "table preview truncated",
    "- extra figure:",
)


def _load_json(path: Path) -> dict:
    return json.loads(path.read_text(encoding="utf-8-sig"))


def _safe_read_text(path: Path) -> str:
    for encoding in ("utf-8", "utf-8-sig"):
        try:
            return path.read_text(encoding=encoding)
        except UnicodeDecodeError:
            continue
    return path.read_text(encoding="utf-8", errors="replace")


def _normalize_text(text: str) -> str:
    cleaned = text or ""
    for wrong, right in MOJIBAKE_REPLACEMENTS.items():
        cleaned = cleaned.replace(wrong, right)
    return cleaned


def _configure_document_styles(document: Document) -> None:
    normal = document.styles["Normal"]
    normal.font.name = "Aptos"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(7)
    normal.paragraph_format.line_spacing = 1.08

    title = document.styles["Title"]
    title.font.name = "Aptos Display"
    title.font.size = Pt(28)
    title.font.bold = False
    title.font.color.rgb = RGBColor(23, 54, 93)
    title.paragraph_format.space_after = Pt(20)

    heading_1 = document.styles["Heading 1"]
    heading_1.font.name = "Aptos Display"
    heading_1.font.size = Pt(20)
    heading_1.font.bold = True
    heading_1.font.color.rgb = RGBColor(23, 54, 93)
    heading_1.paragraph_format.space_before = Pt(20)
    heading_1.paragraph_format.space_after = Pt(8)
    heading_1.paragraph_format.keep_with_next = True

    heading_2 = document.styles["Heading 2"]
    heading_2.font.name = "Aptos"
    heading_2.font.size = Pt(15)
    heading_2.font.bold = True
    heading_2.font.color.rgb = RGBColor(68, 114, 196)
    heading_2.paragraph_format.space_before = Pt(13)
    heading_2.paragraph_format.space_after = Pt(5)
    heading_2.paragraph_format.keep_with_next = True

    heading_3 = document.styles["Heading 3"]
    heading_3.font.name = "Aptos"
    heading_3.font.size = Pt(12.5)
    heading_3.font.bold = True
    heading_3.font.color.rgb = RGBColor(31, 78, 121)
    heading_3.paragraph_format.space_before = Pt(10)
    heading_3.paragraph_format.space_after = Pt(4)
    heading_3.paragraph_format.keep_with_next = True

    heading_4 = document.styles["Heading 4"]
    heading_4.font.name = "Aptos"
    heading_4.font.size = Pt(11)
    heading_4.font.bold = True
    heading_4.font.italic = True
    heading_4.font.color.rgb = RGBColor(68, 114, 196)
    heading_4.paragraph_format.space_before = Pt(7)
    heading_4.paragraph_format.space_after = Pt(3)
    heading_4.paragraph_format.keep_with_next = True


def _humanize_slug(value: str) -> str:
    text = re.sub(r"[_\-]+", " ", value).strip()
    if not text:
        return "Untitled"
    text = re.sub(r"\btest\s*(\d+)\b", lambda m: f"Test {m.group(1)}", text, flags=re.IGNORECASE)
    words = text.split()
    rendered: list[str] = []
    for word in words:
        if re.fullmatch(r"\d+", word):
            rendered.append(word)
        elif len(word) == 1 and word.isalpha():
            rendered.append(word.upper())
        elif word.upper() in {"NAV", "CSV", "JSON", "OOS", "LLM"}:
            rendered.append(word.upper())
        else:
            rendered.append(word.capitalize())
    return " ".join(rendered)


def _friendly_report_label(value: object, *, kind: str = "report") -> str:
    raw = _normalize_text(str(value or "")).strip()
    if not raw:
        return ""
    lowered = raw.lower()
    if "cost_15bps_mainline" in lowered:
        if kind == "bundle":
            return "Mainline Backtest Evidence - 15 bps Cost Assumption"
        if kind == "run":
            return "Mainline backtest run (15 bps cost assumption)"
        return "Mainline backtest report (15 bps cost assumption)"
    if "nightlybatch" in lowered and kind == "bundle":
        return "Latest Backtest Evidence"
    return _humanize_slug(raw)


def _friendly_benchmark_label(value: object) -> str:
    text = _normalize_text(str(value or "")).strip()
    lowered = text.lower()
    if lowered == "universe_ew":
        return "Universe EW"
    if lowered == "static_baseline":
        return "Static baseline"
    if text.upper() == "SPY":
        return "SPY"
    return text.replace("_", " ") if text else ""


def _friendly_regime_label(value: object) -> str:
    text = _normalize_text(str(value or "")).strip()
    lowered = text.lower()
    if lowered == "normal":
        return "Normal regime"
    if lowered == "stress":
        return "Stress regime"
    if lowered == "all":
        return "Full sample"
    return _humanize_slug(text)


def _reader_friendly_item_label(value: str) -> str:
    mappings = {
        "Test 1": "Trading-Cost Robustness",
        "Test 2": "Backtest-Window Robustness",
        "Test 3": "Concentration Robustness",
        "Test 4": "Factor-Weight Robustness",
        "Test 5": "Regime-Trigger Robustness",
        "Test 6": "Drawdown-Control Robustness",
        "Test 7": "Selector-Band Robustness",
        "Test 8": "Implementation-Constraint Robustness",
        "Ablation A": "Component Contribution A",
        "Ablation B": "Component Contribution B",
        "Ablation C": "Component Contribution C",
        "Fixed windows": "Fixed-Window Time Stability",
        "Regime decomposition": "Market-Regime Stability",
        "test_9": "Resampled-Path Robustness",
        "test_10": "Cost-Shock Robustness",
        "test_11": "Factor-Neighbourhood Robustness",
        "test_12": "Rolling Out-of-Sample Check",
        "test_13": "Simulated-Path Return Distribution",
    }
    return mappings.get(value, _humanize_slug(value))


def _friendly_header(value: str) -> str:
    return FRIENDLY_COLUMN_HEADERS.get(value, _humanize_slug(value))


def _specific_asset_label(stem: str) -> str | None:
    lower_stem = (stem or "").lower()
    for key, label in SPECIFIC_ASSET_LABELS.items():
        if key in lower_stem:
            return label
    return None


def _friendly_scenario_value(table_stem: str, value: str) -> str:
    text = (value or "").strip()
    if not text:
        return ""
    stem = (table_stem or "").lower()
    mappings = {
        "test_01": {
            "cost_10bps": "10 bps trading cost",
            "cost_15bps_mainline": "15 bps trading cost (base)",
            "cost_25bps": "25 bps trading cost",
            "cost_40bps": "40 bps trading cost",
        },
        "test_02": {
            "window_minus_6m": "Backtest starts 6 months earlier",
            "window_minus_3m": "Backtest starts 3 months earlier",
            "window_mainline": "Window-test base case",
            "Base Window": "Window-test base case",
            "Base window": "Window-test base case",
            "window_plus_3m": "Backtest starts 3 months later",
            "window_plus_6m": "Backtest starts 6 months later",
        },
        "test_03": {
            "breadth_broader": "Broader selection breadth",
            "breadth_mainline": "Base selection breadth",
            "breadth_tighter": "Tighter selection breadth",
        },
        "test_04": {
            "factor_dividend_down_5pct": "Dividend weight -5 percentage points",
            "factor_dividend_up_5pct": "Dividend weight +5 percentage points",
            "factor_equal_weight": "Equal-weight factor mix",
            "factor_market_technical_down_5pct": "Market-technical weight -5 percentage points",
            "factor_market_technical_up_5pct": "Market-technical weight +5 percentage points",
            "factor_quality_down_5pct": "Quality weight -5 percentage points",
            "factor_quality_up_5pct": "Quality weight +5 percentage points",
            "factor_sentiment_up_5pct": "Sentiment weight +5 percentage points",
            "factor_value_down_5pct": "Value weight -5 percentage points",
            "factor_value_up_5pct": "Value weight +5 percentage points",
        },
        "test_05": {
            "regime_disabled": "Regime trigger disabled",
            "regime_less_sensitive": "Less sensitive regime trigger",
            "regime_mainline": "Base regime trigger",
            "regime_more_sensitive": "More sensitive regime trigger",
        },
        "test_06": {
            "brake_off": "Drawdown brake off",
            "brake_mild": "Mild drawdown brake",
            "brake_mainline": "Base drawdown brake",
            "brake_staircase": "Staircase drawdown brake",
            "brake_aggressive": "Aggressive drawdown brake",
        },
        "test_07": {
            "selector_none": "No selector band",
            "selector_wide": "Wide selector band",
            "selector_medium": "Medium selector band",
            "selector_narrow": "Narrow selector band",
        },
        "test_08": {
            "constraints_none": "No implementation constraint",
            "constraints_weak": "Weak implementation constraint",
            "constraints_medium": "Medium implementation constraint",
            "constraints_strong": "Strong implementation constraint",
        },
        "test_13": {
            "empirical_mean_covariance": "Empirical mean/covariance model",
        },
    }
    for key, mapping in mappings.items():
        if key in stem:
            return mapping.get(text, _humanize_slug(text))
    return _humanize_slug(text)


def _format_table_cell(header: str, value: str, table_stem: str) -> str:
    raw = (value or "").strip()
    if not raw:
        return ""
    if header == "regime":
        return _friendly_regime_label(raw)
    if header == "versus_series":
        return _friendly_benchmark_label(raw)
    if header == "test_key":
        return _reader_friendly_item_label(raw)
    if header in {"scenario_key", "scenario_label"}:
        return _friendly_scenario_value(table_stem, raw)
    if header == "implementation_status":
        return _humanize_slug(raw)
    if header == "title" and "monte carlo path simulation" in raw.lower():
        return "Simulation from realised monthly strategy and benchmark returns"
    numeric = _safe_float(raw)
    if numeric is not None:
        if header in PERCENT_TABLE_HEADERS:
            return f"{numeric:.2f}%"
        if header in PROBABILITY_TABLE_HEADERS:
            return f"{numeric * 100:.1f}%"
        if header in INTEGER_TABLE_HEADERS:
            return f"{numeric:,.0f}"
        return f"{numeric:.3f}"
    return _strip_inline_markdown(raw)


def _friendly_table_rows(rows: list[list[str]], table_stem: str) -> list[list[str]]:
    if not rows:
        return rows
    header = rows[0]
    if "regime" in header and "versus_series" in header:
        regime_idx = header.index("regime")
        versus_idx = header.index("versus_series")
        keep_indices = [idx for idx in range(len(header)) if idx != versus_idx]
        pretty_header = [
            "Regime and benchmark" if idx == regime_idx else _friendly_header(header[idx])
            for idx in keep_indices
        ]
        formatted = [pretty_header]
        for row in rows[1:]:
            padded = row + [""] * (len(header) - len(row))
            combined_label = (
                f"{_friendly_regime_label(padded[regime_idx])} vs "
                f"{_friendly_benchmark_label(padded[versus_idx])}"
            ).strip()
            formatted_row = []
            for idx in keep_indices:
                if idx == regime_idx:
                    formatted_row.append(combined_label)
                else:
                    formatted_row.append(_format_table_cell(header[idx], padded[idx], table_stem))
            formatted.append(formatted_row)
        return formatted
    pretty_header = [_friendly_header(cell) for cell in header]
    formatted = [pretty_header]
    for row in rows[1:]:
        padded = row + [""] * (len(header) - len(row))
        formatted.append(
            [_format_table_cell(header[idx], padded[idx], table_stem) for idx in range(len(header))]
        )
    return formatted


def _preferred_csv_column_indices(header: list[str], table_stem: str) -> list[int]:
    stem = (table_stem or "").lower()
    if "test_13" in stem:
        preferred = [
            "test_key",
            "scenario_key",
            "path_count",
            "sharpe_p50",
            "positive_return_probability",
            "positive_sharpe_probability",
        ]
        indices = [header.index(name) for name in preferred if name in header]
        if len(indices) >= 4:
            return indices
    return []


def _strip_inline_markdown(text: str) -> str:
    cleaned = _normalize_text(text or "")
    cleaned = re.sub(r"`([^`]+)`", r"\1", cleaned)
    cleaned = re.sub(r"\*\*(.+?)\*\*", r"\1", cleaned)
    cleaned = re.sub(r"\*(.+?)\*", r"\1", cleaned)
    return cleaned.strip()


def _looks_like_word_xml(text: str) -> bool:
    sample = (text or "").strip()
    if not sample:
        return False
    lowered = sample.lower()
    return (
        lowered.startswith("<w:")
        or lowered.startswith("</w:")
        or lowered.startswith("&lt;w:")
        or lowered.startswith("w:tbl")
        or lowered.startswith("w:tr")
        or lowered.startswith("w:tc")
        or "<w:" in lowered
        or "&lt;w:" in lowered
    )


def _is_internal_instruction_line(text: str) -> bool:
    line = _strip_inline_markdown(text).strip()
    if not line:
        return False
    if _looks_like_word_xml(line):
        return True
    lower = line.lower()
    return any(lower.startswith(prefix) for prefix in INTERNAL_NOTE_PREFIXES)


def _looks_like_heading(line: str) -> bool:
    compact = line.strip()
    if not compact:
        return False
    if compact.startswith("#"):
        return True
    return bool(re.fullmatch(r"[A-Z][A-Za-z0-9 /\-&:,()]+", compact)) and len(compact.split()) <= 8


def _add_markdown_blocks(document: Document, markdown_text: str) -> None:
    for raw_block in re.split(r"\n\s*\n", markdown_text or ""):
        block = raw_block.strip()
        if not block:
            continue
        if block.startswith("# "):
            document.add_heading(block[2:].strip(), level=1)
            continue
        if block.startswith("## "):
            document.add_heading(block[3:].strip(), level=2)
            continue
        if block.startswith("### "):
            document.add_heading(block[4:].strip(), level=3)
            continue
        lines = []
        for line in block.splitlines():
            if not line.strip():
                continue
            cleaned_line = _strip_inline_markdown(line.strip())
            if not cleaned_line or _is_internal_instruction_line(cleaned_line):
                continue
            lines.append(cleaned_line)
        if not lines:
            continue
        if len(lines) == 1 and _looks_like_heading(lines[0]):
            document.add_heading(lines[0], level=4)
            continue
        if all(line.startswith("- ") for line in lines):
            for line in lines:
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            continue
        for line in lines:
            if line.startswith("- "):
                document.add_paragraph(line[2:].strip(), style="List Bullet")
            else:
                document.add_paragraph(line)


def _clean_note_text(markdown_text: str) -> str:
    cleaned = _strip_inline_markdown(markdown_text or "")
    lines = []
    for raw_line in cleaned.splitlines():
        line = raw_line.strip()
        if not line:
            lines.append("")
            continue
        if _is_internal_instruction_line(line):
            continue
        lines.append(line)
    return "\n".join(lines).strip()


def _has_renderable_markdown_content(markdown_text: str) -> bool:
    for raw_block in re.split(r"\n\s*\n", markdown_text or ""):
        block = raw_block.strip()
        if not block:
            continue
        if block.startswith(("# ", "## ", "### ")):
            return True
        for line in block.splitlines():
            cleaned_line = _strip_inline_markdown(line.strip())
            if cleaned_line and not _is_internal_instruction_line(cleaned_line):
                return True
    return False


def _sanitize_section_body(title: str, body: str) -> str:
    text = _normalize_text(body or "")
    blocked_terms = SECTION_PARAGRAPH_EXCLUSIONS.get(title)
    if not blocked_terms:
        return text
    kept_blocks: list[str] = []
    for raw_block in re.split(r"\n\s*\n", text):
        block = raw_block.strip()
        if not block:
            continue
        if block.lstrip().startswith(("#", "-", "*")):
            kept_lines: list[str] = []
            for line in block.splitlines():
                stripped = line.strip()
                if not stripped:
                    continue
                normalized = _normalize_text(stripped)
                lower_line = normalized.lower()
                if _is_internal_instruction_line(normalized):
                    continue
                if any(term in lower_line for term in blocked_terms):
                    continue
                kept_lines.append(stripped)
            if kept_lines:
                kept_blocks.append("\n".join(kept_lines))
            continue
        sentences = re.split(r"(?<=[.!?])\s+", block)
        kept_sentences = []
        for sentence in sentences:
            cleaned = sentence.strip()
            if not cleaned:
                continue
            if _is_internal_instruction_line(cleaned):
                continue
            lower_sentence = cleaned.lower()
            if any(term in lower_sentence for term in blocked_terms):
                continue
            kept_sentences.append(cleaned)
        if kept_sentences:
            kept_blocks.append(" ".join(kept_sentences))
    return "\n\n".join(kept_blocks)


def _context_snapshot(payload: dict) -> dict:
    snapshot = payload.get("context_snapshot") or {}
    return snapshot if isinstance(snapshot, dict) else {}


def _context_regime_text(payload: dict) -> str:
    snapshot = _context_snapshot(payload)
    for key in ("overview", "market_state", "regime_snapshot", "risk"):
        value = snapshot.get(key)
        if not isinstance(value, dict):
            continue
        for regime_key in ("regime", "current_regime", "state"):
            regime = str(value.get(regime_key) or "").strip()
            if regime:
                return regime
        regime = str(value.get("regime_state") or "").strip()
        if regime:
            return regime
    return ""


def _context_vix_text(payload: dict) -> str:
    snapshot = _context_snapshot(payload)
    for key in ("overview", "market_state", "regime_snapshot", "risk"):
        value = snapshot.get(key)
        if not isinstance(value, dict):
            continue
        for vix_key in ("vix", "vix_level", "current_vix"):
            raw = value.get(vix_key)
            if isinstance(raw, (int, float)):
                return f"{float(raw):.1f}"
            text = str(raw or "").strip()
            if text:
                return text
    return ""


def _performance_summary(payload: dict) -> dict:
    snapshot = _context_snapshot(payload)
    performance = snapshot.get("performance") or {}
    if not isinstance(performance, dict):
        return {}
    summary = performance.get("summary") or {}
    return summary if isinstance(summary, dict) else {}


def _primary_benchmark_summary(payload: dict) -> dict:
    snapshot = _context_snapshot(payload)
    performance = snapshot.get("performance") or {}
    if not isinstance(performance, dict):
        return {}
    comparatives = performance.get("comparatives") or {}
    if not isinstance(comparatives, dict):
        return {}
    primary = comparatives.get("primary_benchmark") or {}
    return primary if isinstance(primary, dict) else {}


def _fmt_pct(value: object) -> str:
    if isinstance(value, (int, float)):
        return f"{float(value):.2f}%"
    return str(value or "").strip()


def _as_float(value: object) -> float | None:
    if isinstance(value, bool):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value or "").strip().replace("%", "")
    if not text:
        return None
    try:
        return float(text)
    except Exception:
        return None


def _previous_update(payload: dict) -> dict:
    previous = _context_snapshot(payload).get("previous_update") or {}
    return previous if isinstance(previous, dict) else {}


def _fmt_change_sentence(
    *,
    label: str,
    current: object,
    previous: object,
    unit: str = "%",
    bigger_is_better: bool | None = None,
) -> str:
    current_num = _as_float(current)
    previous_num = _as_float(previous)
    if current_num is None or previous_num is None:
        return ""
    delta = current_num - previous_num
    if abs(delta) < 0.01:
        return f"{label} is broadly unchanged versus the last materially different data snapshot ({current_num:.2f}{unit})."
    direction = "higher" if delta > 0 else "lower"
    if bigger_is_better is True:
        interpretation = (
            "This is an improvement."
            if delta > 0
            else "This is a softer outcome than the previous snapshot."
        )
    elif bigger_is_better is False:
        interpretation = (
            "This is a more favourable downside profile."
            if delta < 0
            else "This points to a weaker downside profile than before."
        )
    else:
        interpretation = ""
    sentence = (
        f"{label} moved from {previous_num:.2f}{unit} to {current_num:.2f}{unit}, "
        f"which is {direction} by {abs(delta):.2f}{unit}."
    )
    if interpretation:
        sentence += f" {interpretation}"
    return sentence


def _fallback_section_body(title: str, payload: dict) -> str:
    aliased_title = SECTION_FALLBACK_ALIASES.get(title)
    if aliased_title:
        return _fallback_section_body(aliased_title, payload)
    regime = _context_regime_text(payload)
    vix_text = _context_vix_text(payload)
    perf = _performance_summary(payload)
    primary = _primary_benchmark_summary(payload)
    previous_update = _previous_update(payload)
    if title == "Strategy And Portfolio Construction":
        return (
            "The strategy is a systematic long-only US equity portfolio that ranks eligible stocks using diversified factor evidence rather than a single signal. "
            "The process combines quality, value, market/technical, and dividend information, then applies regime-aware tilts and covariance-aware construction so that expected return, diversification, and risk concentration are considered together. "
            "The portfolio is intended to add value versus SPY as the final investor-facing baseline, with the same-universe equal-weight benchmark retained as a supporting stock-selection check, while remaining implementable through explicit turnover, single-name, and sector constraints."
        )
    if title == "Backtest Design":
        benchmark_label = str(primary.get("label") or "SPY").strip()
        is_universe_internal = benchmark_label.lower().replace(" ", "_") in {
            "universe_ew",
            "universe_equal_weight",
            "equal_weight_universe",
        }
        benchmark_sentence = (
            "The backtest uses the same-universe equal-weight series as an internal active-management comparator, while the final investor-facing baseline is SPY."
            if is_universe_internal
            else f"The backtest is interpreted primarily against {benchmark_label}, with the same-universe equal-weight series retained as a supporting internal stock-selection comparator."
        )
        return (
            f"{benchmark_sentence} "
            "Signals are intended to be point-in-time, trades are evaluated after the rebalance decision rather than before it, and returns are assessed net of explicit transaction costs. "
            "This design makes the results more useful for an investor because it separates market exposure, active construction value, implementation cost, and downside behaviour."
        )
    if title == "Product Snapshot":
        annualized = _fmt_pct(perf.get("annualized_return_pct"))
        excess = _fmt_pct(primary.get("excess_return_annualized_pct"))
        drawdown = _fmt_pct(perf.get("max_drawdown_pct"))
        if regime and vix_text and annualized and excess:
            regime_descriptor = "pressured" if regime.lower() == "stress" else "more supportive"
            return (
                f"The product is currently operating in a {regime} market setting with VIX around {vix_text}. "
                f"Even in this environment, the latest history still points to annualized return around {annualized} and benchmark-relative excess return around {excess}, so the current picture is {regime_descriptor} but still competitive rather than broken. "
                f"The key thing users should keep in mind is that downside tolerance still matters, with maximum drawdown around {drawdown}."
            )
        if regime and vix_text:
            return (
                f"The product is currently operating in a {regime} market setting with VIX around {vix_text}. "
                "Users should read the latest evidence as a current-state assessment of whether this environment is supportive or pressured, rather than as a promise that the same conditions will persist."
            )
        if regime:
            return (
                f"The product is currently operating in a {regime} market setting. "
                "The main question for users is whether that environment still supports the expected edge of the strategy."
            )
        return "This section summarises the product's current state: the type of market setting it is facing now, whether the recent picture looks supportive or pressured, and the single most important point a user should understand at this stage."
    if title == "Robustness Watch":
        risk = _context_snapshot(payload).get("risk") or {}
        normal_excess = ""
        stress_excess = ""
        if isinstance(risk, dict):
            summary = risk.get("summary") or {}
            if isinstance(summary, dict):
                normal_excess = _fmt_pct(summary.get("normal_excess_pct"))
                stress_excess = _fmt_pct(summary.get("stress_excess_pct"))
        if normal_excess and stress_excess:
            return (
                f"The latest validation evidence broadly supports the current product message: in calmer conditions the strategy has shown stronger excess return around {normal_excess}, while in stress conditions it has still preserved positive benchmark-relative excess return around {stress_excess}. "
                "Users should therefore read the robustness work as evidence that the edge looks more like a moderated advantage under pressure than a result that disappears as soon as assumptions move."
            )
        return (
            "The latest validation evidence should be read as a credibility check on the current product message. "
            "The practical question for users is whether modest implementation changes, different market slices, and uncertainty scenarios leave the main interpretation broadly intact or expose a weak point that deserves caution."
        )
    if title == "Performance Update":
        annualized = _fmt_pct(perf.get("annualized_return_pct"))
        benchmark_label = str(primary.get("label") or "the primary benchmark").strip()
        benchmark_excess = _fmt_pct(primary.get("excess_return_annualized_pct"))
        sharpe = perf.get("rolling_sharpe") or perf.get("mainline_realized_sharpe")
        drawdown = _fmt_pct(perf.get("max_drawdown_pct"))
        volatility = _fmt_pct(perf.get("volatility_pct"))
        parts: list[str] = []
        if annualized and benchmark_excess:
            parts.append(
                f"The latest history still points to annualized return around {annualized}, with benchmark-relative excess return around {benchmark_excess} versus {benchmark_label}."
            )
        elif annualized:
            parts.append(
                f"The latest history still points to annualized return around {annualized}, which is the main headline result users should focus on."
            )
        if sharpe not in (None, "") and drawdown:
            parts.append(
                f"Risk-adjusted quality remains reasonable, with Sharpe around {float(sharpe):.2f} and maximum drawdown around {drawdown}."
            )
        elif drawdown:
            parts.append(
                f"The key downside statistic remains maximum drawdown around {drawdown}, which defines the level of pain users would have needed to tolerate historically."
            )
        if volatility:
            parts.append(
                f"Volatility remains around {volatility}, which means the product is not behaving like a low-volatility substitute and should still be read as an equity-risk strategy."
            )
        return " ".join(parts) or (
            "This section should focus on the few numbers that best summarise current performance, how that compares with the benchmark, and how much downside risk users would have needed to tolerate."
        )
    if title == "What Changed Since The Last Update":
        if not previous_update:
            return ""
        previous_perf = previous_update.get("performance_summary") or {}
        previous_risk = previous_update.get("risk_summary") or {}
        previous_robustness = previous_update.get("robustness_summary") or {}
        previous_report_id = str(previous_update.get("report_id") or "").strip()
        changes: list[str] = []
        annualized_sentence = _fmt_change_sentence(
            label="Annualized return",
            current=perf.get("annualized_return_pct"),
            previous=previous_perf.get("annualized_return_pct"),
            bigger_is_better=True,
        )
        if annualized_sentence:
            changes.append(annualized_sentence)
        excess_sentence = _fmt_change_sentence(
            label="Benchmark-relative excess return",
            current=primary.get("excess_return_annualized_pct", perf.get("excess_return_pct")),
            previous=previous_perf.get("excess_return_pct"),
            bigger_is_better=True,
        )
        if excess_sentence:
            changes.append(excess_sentence)
        drawdown_sentence = _fmt_change_sentence(
            label="Maximum drawdown",
            current=perf.get("max_drawdown_pct"),
            previous=previous_perf.get("max_drawdown_pct"),
            bigger_is_better=False,
        )
        if drawdown_sentence:
            changes.append(drawdown_sentence)
        current_regime = regime
        previous_regime = str(
            previous_risk.get("regime") or previous_risk.get("current_regime") or ""
        ).strip()
        if current_regime and previous_regime and current_regime.lower() != previous_regime.lower():
            changes.append(
                f"The market backdrop has shifted from {previous_regime} to {current_regime}, so users should expect the product experience to feel different even if the long-run investment case is unchanged."
            )
        current_sharpe = (
            (_context_snapshot(payload).get("robustness") or {}).get("mainline_realized") or {}
        ).get("rolling_sharpe")
        previous_sharpe = (previous_robustness.get("mainline_realized") or {}).get("rolling_sharpe")
        sharpe_sentence = _fmt_change_sentence(
            label="Risk-adjusted quality (Sharpe ratio)",
            current=current_sharpe,
            previous=previous_sharpe,
            unit="",
            bigger_is_better=True,
        )
        if sharpe_sentence:
            changes.append(sharpe_sentence)
        if not changes:
            baseline_text = "The latest run looks broadly similar to the previous materially different data snapshot"
            if previous_report_id:
                baseline_text += f" ({previous_report_id})"
            baseline_text += ", so this update should be read as confirmation of the current product message rather than a material shift."
            return baseline_text
        lead = "Compared with the last materially different data snapshot"
        if previous_report_id:
            lead += f" ({previous_report_id})"
        lead += ", the main changes are as follows:"
        return "\n\n".join([lead, *changes[:4]])
    if title == "Risk And Regime Update":
        drawdown = _fmt_pct(perf.get("max_drawdown_pct"))
        volatility = _fmt_pct(perf.get("volatility_pct"))
        excess = _fmt_pct(primary.get("excess_return_annualized_pct"))
        parts: list[str] = []
        if regime and vix_text:
            if regime.lower() == "stress":
                parts.append(
                    f"The product is currently operating in a {regime} environment with VIX around {vix_text}, which should be read as a more demanding backdrop for absolute returns and user comfort."
                )
            else:
                parts.append(
                    f"The product is currently operating in a {regime} environment with VIX around {vix_text}, which is a calmer backdrop than a stress phase but still one that needs to justify the strategy's edge through realised results."
                )
        elif regime:
            if regime.lower() == "stress":
                parts.append(
                    f"The product is currently operating in a {regime} environment, which means recent behaviour should be judged against a more difficult market backdrop than a calm equity period."
                )
            else:
                parts.append(
                    f"The product is currently operating in a {regime} environment, so recent behaviour should be judged as a more standard market state rather than as a crisis-style stress episode."
                )
        if drawdown and volatility:
            parts.append(
                f"Current risk should therefore be understood through two lenses: historical drawdown tolerance around {drawdown} and realised volatility around {volatility}, both of which imply that this remains an equity-risk product rather than a defensive cash-like holding."
            )
        elif drawdown:
            parts.append(
                f"The most important current risk signal remains drawdown tolerance around {drawdown}, which defines how uncomfortable a difficult holding period could become."
            )
        if excess:
            parts.append(
                f"Even so, benchmark-relative excess return remains around {excess}, so the current picture is better described as pressured but still competitively positioned rather than structurally broken."
            )
        return " ".join(parts) or (
            "This section should explain the current market setting and the main risk pressures in plain terms, with an emphasis on what those pressures mean for user experience rather than on technical factor details."
        )
    if title == "What Users Should Watch Next":
        lead = "Users should keep the next review focused on a small set of practical signals."
        if regime and vix_text:
            lead = f"Users should keep the next review focused on a small set of practical signals while the product operates in a {regime} regime with VIX around {vix_text}."
        elif regime:
            lead = f"Users should keep the next review focused on a small set of practical signals while the product operates in a {regime} regime."
        return "\n\n".join(
            [
                lead,
                "First, watch whether benchmark-relative excess return remains positive, because persistent outperformance versus the reference portfolios is the clearest sign that the strategy is still adding value.",
                "Second, watch whether drawdown behaviour deepens if the current regime persists, because a strategy can remain relatively strong versus a benchmark while still becoming harder to hold through prolonged stress.",
                "Third, watch whether implementation friction such as turnover and trading cost starts to erode the edge, because that would weaken confidence in real-world usability even if the headline backtest numbers still look respectable.",
            ]
        )
    return ""


def _section_needs_fallback(title: str, body: str) -> bool:
    title = SECTION_FALLBACK_ALIASES.get(title, title)
    cleaned = _normalize_text(body or "").strip()
    if not cleaned:
        return True
    if not _has_renderable_markdown_content(cleaned):
        return True
    low = cleaned.lower()
    if title == "Product Snapshot":
        forbidden = [
            "data health",
            "dag",
            "pipeline",
            "artifacts",
            "artifact",
            "handoff pack",
            "timestamped",
            "delivered",
            "delivery",
            "checks passing",
            "run id:",
        ]
        return (
            len(cleaned.split()) < 25
            or any(term in low for term in forbidden)
            or not any(
                phrase in low
                for phrase in ["market", "environment", "supportive", "pressured", "product"]
            )
        )
    if title == "Performance Update":
        metric_markers = [
            "annualized return",
            "cumulative return",
            "excess return",
            "information ratio",
            "tracking error",
            "sharpe",
            "sortino",
            "volatility",
            "drawdown",
        ]
        metric_count = sum(1 for term in metric_markers if term in low)
        return len(cleaned.split()) < 35 or metric_count >= 7
    if title == "Robustness Watch":
        forbidden = [
            "completed",
            "acceptance matrix",
            "remaining block",
            "missing rows",
            "partial implementation",
            "deterministic tests",
            "stochastic tests",
            "ablation blocks",
            "partial block",
            "passed",
        ]
        return (
            len(cleaned.split()) < 30
            or any(term in low for term in forbidden)
            or not any(
                phrase in low
                for phrase in ["stable", "fragile", "confidence", "trust", "uncertainty"]
            )
        )
    if title == "Risk And Regime Update":
        forbidden = [
            "quality (",
            "value (",
            "market technical",
            "dividend tilt",
            "momentum shift",
            "factor exposures",
            "unchanged factor exposures",
        ]
        return (
            len(cleaned.split()) < 35
            or any(term in low for term in forbidden)
            or not any(
                phrase in low
                for phrase in [
                    "market",
                    "regime",
                    "risk",
                    "pressure",
                    "environment",
                    "volatility",
                    "drawdown",
                ]
            )
        )
    if title == "What Changed Since The Last Update":
        forbidden = [
            "first detailed robustness",
            "no prior update data",
            "acceptance matrix",
            "blocks completed",
            "tests completed",
            "partial block",
            "partial implementation",
            "fixed window subperiod tables",
        ]
        return len(cleaned.split()) < 20 or any(term in low for term in forbidden)
    if title == "What Users Should Watch Next":
        forbidden = [
            "future updates should",
            "follow-up",
            "needs resolution",
            "team",
            "internal",
            "core performance evidence",
            "key robustness evidence",
            "deterministic tests",
            "stochastic tests",
            "partial acceptance",
            "partial block",
            "acceptance block",
            "fixed window subperiod tables",
            "recommend completing",
            "outstanding",
        ]
        return (
            len(cleaned.split()) < 45
            or any(term in low for term in forbidden)
            or not any(
                phrase in low
                for phrase in [
                    "watch",
                    "monitor",
                    "confidence",
                    "warning",
                    "strengthen",
                    "weaken",
                    "signal",
                ]
            )
        )
    return False


def _sanitize_evidence_analysis(text: str) -> str:
    body = _normalize_text(text or "")
    body = body.replace(
        "confirming robustness to random market fluctuations and cost perturbations.",
        "supporting robustness under simulated return-path variation.",
    )
    replacements = {
        r"\bTest\s*2[’']s sensitivity analysis": "The backtest-window robustness analysis",
        r"\bTest\s*13\b": "The simulated-path check",
        r"\btest\s*13\b": "the simulated-path check",
        r"\bdeterministic robustness check\b": "parameter-robustness check",
        r"\bstochastic tests\b": "simulation-based checks",
        r"\bstochastic test\b": "simulation-based check",
        r"\bpersistent alpha generation\b": "historical active-return evidence",
        r"\bstable, repeatable performance\b": "historically stable performance under the tested assumptions",
        r"\bnot artifacts of historical data\b": "not solely an artifact of one realised historical path",
        r"\blikely to persist under varied future market conditions\b": "not solely dependent on one realised path, although future persistence is not guaranteed",
        r"\bstrengthens confidence in the strategy[â€™']s reliability\b": "supports the reliability of the historical result",
        r"\benhancing confidence in its generalizability and reducing timing risk\b": "reducing concern that the result depends on one entry point",
        r"\bInvestors gain assurance from this evidence that\b": "Investors can read this evidence as showing that",
        r"\bnot solely an artifact of one realised historical path but are not solely dependent on one realised path, although future persistence is not guaranteed\b": "not solely an artifact of one realised historical path, although future persistence is not guaranteed",
        r"\bThis supports confidence in the strategy[â€™']s durability and suitability for inclusion in diversified portfolios seeking consistent alpha\.\b": "This makes the simulation useful as a risk check rather than as a standalone allocation recommendation.",
    }
    for pattern, replacement in replacements.items():
        body = re.sub(pattern, replacement, body, flags=re.IGNORECASE)
    body = re.sub(
        r"\bstrengthens confidence in the strategy.?s reliability\b",
        "supports the reliability of the historical result",
        body,
        flags=re.IGNORECASE,
    )
    body = re.sub(
        r"\bThis supports confidence in the strategy.?s durability and suitability for inclusion in diversified portfolios seeking consistent alpha\.",
        "This makes the simulation useful as a risk check rather than as a standalone allocation recommendation.",
        body,
        flags=re.IGNORECASE,
    )
    blocked_terms = [
        "data health",
        "dag health",
        "acceptance matrix",
        "robustness blocks passed",
        "partial robustness block",
        "partial acceptance",
        "remaining block",
        "future updates should",
        "follow-up",
        "needs resolution",
        "should be completed",
        "research team",
        "internal next step",
        "future work",
        "team task",
    ]
    kept_sentences: list[str] = []
    for sentence in re.split(r"(?<=[.!?])\s+", body):
        cleaned = sentence.strip()
        if not cleaned:
            continue
        lower_sentence = cleaned.lower()
        if any(term in lower_sentence for term in blocked_terms):
            continue
        kept_sentences.append(cleaned)
    return " ".join(kept_sentences).strip()


def _split_evidence_analysis_text(text: str) -> tuple[str, str]:
    body = _normalize_text(text or "")
    if not body:
        return "", ""
    investor_match = re.search(r"(?is)\bInvestor interpretation:\s*", body)
    if not investor_match:
        analysis = re.sub(r"(?im)^Analysis:\s*", "", body).strip()
        return analysis, ""
    analysis = body[: investor_match.start()].strip()
    investor_interpretation = body[investor_match.end() :].strip()
    analysis = re.sub(r"(?im)^Analysis:\s*", "", analysis).strip()
    investor_interpretation = re.sub(
        r"(?im)^Investor interpretation:\s*",
        "",
        investor_interpretation,
    ).strip()
    return analysis, investor_interpretation


def _clean_investor_interpretation(text: str) -> str:
    body = _sanitize_evidence_analysis(text or "")
    body = re.sub(r"(?im)^Investor interpretation:\s*", "", body).strip()
    if len(body.split()) < 8:
        return ""
    blocked_phrases = [
        "use this figure",
        "use this table",
        "include this",
        "this section should",
        "the report should include",
        "the writer should",
    ]
    if any(phrase in body.lower() for phrase in blocked_phrases):
        return ""
    return body


def _add_investor_interpretation(document: Document, text: str) -> None:
    clean_text = _clean_investor_interpretation(text)
    if not clean_text:
        return
    paragraph = document.add_paragraph()
    paragraph.add_run("Investor interpretation: ").bold = True
    paragraph.add_run(clean_text)


def _fallback_investor_interpretation(
    title: str,
    asset_names: list[str],
    *,
    part_label: str = "",
    report_dir: Path | None = None,
) -> str:
    block_key = _asset_key_from_block(title, asset_names)
    if report_dir is not None:
        for key, interpretation in MAIN_REPORT_INVESTOR_INTERPRETATIONS.items():
            if key in block_key:
                return interpretation
    if part_label in ROBUSTNESS_INVESTOR_INTERPRETATIONS:
        return ROBUSTNESS_INVESTOR_INTERPRETATIONS[part_label]
    lowered = f"{block_key} {part_label}".lower()
    if "regime" in lowered:
        return (
            "For investors, this evidence is most useful as a condition map: it explains when the strategy is likely to feel strong or weak, "
            "so confidence should be tied to the market environment rather than to the full-sample result alone."
        )
    if "cost" in lowered or "turnover" in lowered:
        return (
            "For investors, this evidence turns implementation from an assumption into a monitoring item. The strategy is more convincing when the return edge "
            "remains visible after plausible cost and turnover pressure."
        )
    if "window" in lowered or "oos" in lowered or "sample" in lowered:
        return (
            "For investors, this evidence reduces reliance on a single favourable sample. It supports confidence only to the extent that the conclusion remains "
            "recognisable across alternative windows or validation paths."
        )
    return (
        "For investors, this evidence should be read as a confidence check around the main backtest conclusion. It is useful when it confirms the investment case, "
        "and it becomes a monitoring signal if it points to fragility or concentration."
    )


def _block_investor_interpretation(
    block: dict[str, object],
    title: str,
    asset_names: list[str],
    *,
    inline_interpretation: str = "",
    part_label: str = "",
    report_dir: Path | None = None,
) -> str:
    block_key = _asset_key_from_block(title, asset_names)
    if report_dir is not None and any(
        key in block_key for key in MAIN_REPORT_INVESTOR_OVERRIDE_KEYS
    ):
        return _clean_investor_interpretation(
            _fallback_investor_interpretation(
                title,
                asset_names,
                part_label=part_label,
                report_dir=report_dir,
            )
        )
    for candidate in (
        str(block.get("investor_interpretation") or ""),
        inline_interpretation,
        _fallback_investor_interpretation(
            title,
            asset_names,
            part_label=part_label,
            report_dir=report_dir,
        ),
    ):
        clean_text = _clean_investor_interpretation(candidate)
        if clean_text:
            return clean_text
    return ""


def _plain_language_evidence_note(title: str) -> str | None:
    normalized = _normalize_text(title or "").strip()
    if not normalized:
        return None
    if normalized in PLAIN_LANGUAGE_EVIDENCE_NOTES:
        return PLAIN_LANGUAGE_EVIDENCE_NOTES[normalized]
    lowered = normalized.lower()
    for key, note in PLAIN_LANGUAGE_EVIDENCE_NOTES.items():
        if lowered == key.lower():
            return note
    return None


def _populate_word_table(table, rows: list[list[str]]) -> None:
    for row_idx, row in enumerate(rows):
        for col_idx, value in enumerate(row):
            table.rows[row_idx].cells[col_idx].text = _table_cell_text(value)


def _table_cell_text(value: object) -> str:
    text = _normalize_text(str(value or "")).strip()
    if len(text) > 32:
        text = text.replace("_", " ")
        text = text.replace("\\", " / ").replace("/", " / ")
    return re.sub(r"\s+", " ", text).strip()


def _rows_need_flow_layout(rows: list[list[str]]) -> bool:
    if not rows:
        return False
    column_count = max(len(row) for row in rows)
    if column_count > 4:
        return True
    return any(
        len(str(cell or "")) > 42 and any(marker in str(cell) for marker in ("_", "\\", "/"))
        for row in rows
        for cell in row
    )


def _add_flow_rows(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    headers = [_table_cell_text(cell) for cell in rows[0]]
    body_rows = rows[1:] if len(rows) > 1 else []
    if len(headers) == 2:
        for row in body_rows:
            if len(row) < 2:
                continue
            label = _table_cell_text(row[0])
            value = _table_cell_text(row[1])
            if not label or not value:
                continue
            paragraph = document.add_paragraph(style="List Bullet")
            paragraph.add_run(f"{label}: ").bold = True
            paragraph.add_run(value)
        return
    for row in body_rows:
        normalized_row = [_table_cell_text(cell) for cell in row]
        if not any(normalized_row):
            continue
        label = normalized_row[0] if normalized_row else "Evidence row"
        details = []
        for header, value in zip(headers[1:], normalized_row[1:]):
            if header and value:
                details.append(f"{header}: {value}")
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.add_run(f"{label}: ").bold = True
        paragraph.add_run("; ".join(details) if details else "No additional values shown.")


def _add_csv_table_chunk(document: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    if _rows_need_flow_layout(rows):
        _add_flow_rows(document, rows)
        return
    table = document.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Table Grid"
    table.autofit = False
    usable_width = 6.2
    col_width = Inches(usable_width / max(1, len(rows[0])))
    for column in table.columns:
        for cell in column.cells:
            cell.width = col_width
    _populate_word_table(table, rows)


def _add_csv_table(
    document: Document,
    csv_path: Path,
    max_rows: int = MAX_WORD_TABLE_ROWS,
    *,
    compact: bool = False,
) -> None:
    rows: list[list[str]] = []
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        reader = csv.reader(handle)
        for idx, row in enumerate(reader):
            rows.append([str(cell) for cell in row])
            if idx >= max_rows:
                break
    if not rows:
        document.add_paragraph("Table file is empty.")
        return
    width = max(len(row) for row in rows)
    normalized = [row + [""] * (width - len(row)) for row in rows]
    preferred_indices = _preferred_csv_column_indices(normalized[0], csv_path.stem)
    if preferred_indices:
        normalized = [[row[idx] for idx in preferred_indices] for row in normalized]
    normalized = _friendly_table_rows(normalized, csv_path.stem)
    row_limit = max_rows if compact else min(max_rows, MAX_WORD_TABLE_ROWS)
    trimmed = [row[:MAX_WORD_TABLE_COLUMNS] for row in normalized[:row_limit]]
    _add_csv_table_chunk(document, trimmed)


def _csv_has_meaningful_content(csv_path: Path, sample_rows: int = 40) -> bool:
    try:
        with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            rows = list(reader)
    except Exception:
        return False
    if len(rows) <= 1:
        return False
    for row in rows[1:sample_rows]:
        for cell in row:
            if (cell or "").strip():
                return True
    return False


def _add_json_summary(document: Document, json_path: Path) -> None:
    payload = _load_json(json_path)
    if not isinstance(payload, dict):
        document.add_paragraph("JSON payload is not an object; refer to the source file.")
        return
    preferred_keys = [
        "report_name",
        "run_name",
        "start_date",
        "end_date",
        "rebalance_frequency",
        "benchmark_ticker",
        "primary_benchmark",
        "annualized_return",
        "excess_return_vs_primary",
        "sharpe_ratio",
        "max_drawdown",
        "annualized_volatility",
        "avg_monthly_turnover",
        "annualized_turnover_ratio",
        "transaction_cost_bps",
        "scorecard_passed",
        "scorecard_total",
    ]
    summary_rows: list[list[str]] = [["Field", "Value"]]
    for key in preferred_keys:
        if key not in payload:
            continue
        value = payload.get(key)
        if isinstance(value, (dict, list)):
            continue
        rendered = _format_json_value(key, value)
        if not rendered or _looks_like_word_xml(rendered):
            continue
        summary_rows.append([_friendly_header(key), rendered])
    if len(summary_rows) == 1:
        for key, value in payload.items():
            if isinstance(value, (dict, list)):
                continue
            rendered = _format_json_value(str(key), value)
            if not rendered or _looks_like_word_xml(rendered):
                continue
            summary_rows.append([_friendly_header(str(key)), rendered])
            if len(summary_rows) >= 9:
                break
    if len(summary_rows) == 1:
        document.add_paragraph("Structured summary is unavailable for this JSON payload.")
        return
    _add_flow_rows(document, summary_rows)


def _json_summary_lines(json_path: Path) -> list[str]:
    payload = _load_json(json_path)
    if not isinstance(payload, dict):
        return []
    preferred_keys = [
        "report_name",
        "run_name",
        "start_date",
        "end_date",
        "rebalance_frequency",
        "primary_benchmark",
        "annualized_return",
        "excess_return_vs_primary",
        "sharpe_ratio",
        "max_drawdown",
        "annualized_volatility",
        "avg_monthly_turnover",
        "annualized_turnover_ratio",
        "transaction_cost_bps",
    ]
    lines: list[str] = []
    for key in preferred_keys:
        if key not in payload:
            continue
        value = payload.get(key)
        if isinstance(value, (dict, list)):
            continue
        rendered = _format_json_value(key, value)
        if not rendered or _looks_like_word_xml(rendered):
            continue
        lines.append(f"{_friendly_header(key)}: {rendered}")
    return lines


def _format_json_value(key: str, value: object) -> str:
    if value is None:
        return ""
    if isinstance(value, bool):
        return "Yes" if value else "No"
    text_key = str(key)
    if isinstance(value, (int, float)) and not isinstance(value, bool):
        numeric = float(value)
        if text_key in {
            "total_return",
            "annualized_return",
            "excess_return_vs_primary",
            "max_drawdown",
            "annualized_volatility",
            "avg_monthly_turnover",
            "annualized_turnover_ratio",
            "avg_executed_turnover",
        }:
            return f"{numeric:.2f}%"
        if text_key in {"transaction_cost_bps"}:
            return f"{numeric:.0f} bps"
        if text_key in {"sharpe_ratio"}:
            return f"{numeric:.2f}"
        return f"{numeric:.2f}".rstrip("0").rstrip(".")
    if text_key == "report_name":
        return _friendly_report_label(value, kind="report")
    if text_key == "run_name":
        return _friendly_report_label(value, kind="run")
    if text_key in {"primary_benchmark", "secondary_benchmark", "benchmark_ticker"}:
        return _friendly_benchmark_label(value)
    rendered = _normalize_text(str(value)).strip()
    return "" if _looks_like_word_xml(rendered) else rendered


def _add_metric_snapshot_as_text(
    document: Document, blocks: list[tuple[str, list[tuple[str, str]]]]
) -> None:
    heading = document.add_paragraph()
    heading.paragraph_format.space_before = Pt(4)
    heading.paragraph_format.space_after = Pt(2)
    heading_run = heading.add_run("Key Metrics Snapshot")
    heading_run.bold = True
    heading_run.font.color.rgb = RGBColor(31, 78, 121)
    heading_run.font.size = Pt(11)

    for title, rows in blocks:
        para = document.add_paragraph()
        para.paragraph_format.left_indent = Pt(10)
        para.paragraph_format.first_line_indent = Pt(-10)
        para.paragraph_format.space_before = Pt(0)
        para.paragraph_format.space_after = Pt(0.8)
        title_run = para.add_run(f"{title}: ")
        title_run.bold = True
        title_run.font.color.rgb = RGBColor(31, 78, 121)
        title_run.font.size = Pt(9.8)
        value_run = para.add_run("; ".join(f"{label} {value or '-'}" for label, value in rows))
        value_run.font.size = Pt(9.8)


def _acceptance_summary_from_evidence(evidence_dir: Path | None) -> tuple[str, str]:
    if evidence_dir is None:
        return "", ""
    matrix_path = evidence_dir / "part_5_dashboard_and_conclusions" / "acceptance_matrix.csv"
    if not matrix_path.exists():
        return "", ""
    try:
        with matrix_path.open("r", encoding="utf-8-sig", newline="") as handle:
            rows = list(csv.DictReader(handle))
    except OSError:
        return "", ""
    if not rows:
        return "", ""
    total = len(rows)
    completed = sum(1 for row in rows if (row.get("status") or "").strip().lower() == "completed")
    partial = next(
        (row for row in rows if (row.get("status") or "").strip().lower() != "completed"), None
    )
    partial_label = ""
    if partial:
        partial_label = _normalize_text(
            str(partial.get("label") or partial.get("item_key") or "one acceptance item")
        )
    return f"{completed}/{total}", partial_label


def _validation_scope_note(scorecard: str, acceptance: str, partial_label: str) -> str:
    if not acceptance:
        return ""
    if partial_label:
        return (
            "Validation note: the main scorecard and robustness acceptance count use different scopes. "
            f"The robustness evidence pack records {acceptance} completed acceptance items, with {partial_label} still partial."
        )
    if scorecard and scorecard != acceptance:
        return (
            "Validation note: the main scorecard and robustness acceptance count use different scopes, "
            f"so {scorecard} and {acceptance} should not be read as duplicate measures."
        )
    return ""


def _add_opening_metric_snapshot(
    document: Document, summary_path: Path, evidence_dir: Path | None = None
) -> None:
    payload = _load_json(summary_path)
    if not isinstance(payload, dict):
        _add_json_summary(document, summary_path)
        return

    benchmark = _friendly_benchmark_label(payload.get("primary_benchmark") or "SPY")
    secondary = _friendly_benchmark_label(
        payload.get("secondary_benchmark") or payload.get("benchmark_ticker") or "SPY"
    )
    start_date = _format_json_value("start_date", payload.get("start_date"))
    end_date = _format_json_value("end_date", payload.get("end_date"))
    rebalance = _format_json_value("rebalance_frequency", payload.get("rebalance_frequency"))
    cost = _format_json_value("transaction_cost_bps", payload.get("transaction_cost_bps"))
    document.add_paragraph(
        f"Mainline run: {start_date} to {end_date}; {rebalance} rebalance; {cost} transaction-cost assumption; benchmarks: {benchmark} and {secondary}."
    )

    scorecard_passed = payload.get("scorecard_passed")
    scorecard_total = payload.get("scorecard_total")
    scorecard = (
        f"{int(scorecard_passed)}/{int(scorecard_total)}"
        if isinstance(scorecard_passed, (int, float)) and isinstance(scorecard_total, (int, float))
        else "-"
    )
    acceptance, partial_label = _acceptance_summary_from_evidence(evidence_dir)
    blocks = [
        (
            "Return",
            [
                (
                    "Annualized return",
                    _format_json_value("annualized_return", payload.get("annualized_return")),
                ),
                (
                    f"Excess vs {benchmark}",
                    _format_json_value(
                        "excess_return_vs_primary", payload.get("excess_return_vs_primary")
                    ),
                ),
                ("Total return", _format_json_value("total_return", payload.get("total_return"))),
            ],
        ),
        (
            "Risk",
            [
                ("Sharpe ratio", _format_json_value("sharpe_ratio", payload.get("sharpe_ratio"))),
                ("Max drawdown", _format_json_value("max_drawdown", payload.get("max_drawdown"))),
                (
                    "Volatility",
                    _format_json_value(
                        "annualized_volatility", payload.get("annualized_volatility")
                    ),
                ),
            ],
        ),
        (
            "Implementation",
            [
                (
                    "Avg monthly turnover",
                    _format_json_value("avg_monthly_turnover", payload.get("avg_monthly_turnover")),
                ),
                (
                    "Annualized turnover",
                    _format_json_value(
                        "annualized_turnover_ratio", payload.get("annualized_turnover_ratio")
                    ),
                ),
                ("Transaction cost", cost),
            ],
        ),
        (
            "Validation",
            [
                ("Main scorecard", scorecard),
                ("Robustness acceptance", acceptance or "-"),
                ("Periods", _format_json_value("periods", payload.get("periods"))),
                ("Charts", _format_json_value("chart_count", payload.get("chart_count"))),
            ],
        ),
    ]
    _add_metric_snapshot_as_text(document, blocks)
    note = _validation_scope_note(scorecard, acceptance, partial_label)
    if note:
        note_para = document.add_paragraph(note)
        note_para.paragraph_format.space_before = Pt(2)
        note_para.paragraph_format.space_after = Pt(2)
        note_para.runs[0].font.size = Pt(9.2)


def _opening_takeaway(summary_path: Path) -> str:
    payload = _load_json(summary_path)
    if not isinstance(payload, dict):
        return "Takeaway: the report should be read as a data-backed investment view, with performance evidence interpreted alongside risk, implementation, and robustness checks."
    benchmark = _friendly_benchmark_label(payload.get("primary_benchmark") or "SPY")
    annual_return = _format_json_value("annualized_return", payload.get("annualized_return"))
    excess_return = _format_json_value(
        "excess_return_vs_primary", payload.get("excess_return_vs_primary")
    )
    sharpe = _format_json_value("sharpe_ratio", payload.get("sharpe_ratio"))
    drawdown = _format_json_value("max_drawdown", payload.get("max_drawdown"))
    return (
        f"Takeaway: the strategy shows a positive active-return case"
        f"{f' with {annual_return} annualized return and {excess_return} excess return versus {benchmark}' if annual_return and excess_return else ''}, "
        f"but the conclusion should be read alongside its equity-like risk profile"
        f"{f' (Sharpe {sharpe}, max drawdown {drawdown})' if sharpe and drawdown else ''}."
    )


def _guess_image_note(filename: str) -> str:
    lower_name = filename.lower()
    for keyword, note in IMAGE_NOTE_KEYWORDS.items():
        if keyword in lower_name:
            return note
    return "This visual is included as supporting evidence for the written explanation in this section."


def _part_summary(part_label: str) -> str:
    for key, summary in PART_SUMMARIES.items():
        if part_label.startswith(key):
            return summary
    return "This section contains robustness evidence grouped to support the written report."


def _part_report_use(part_label: str) -> str:
    for key, summary in PART_REPORT_USE.items():
        if part_label.startswith(key):
            return summary
    return "Use this evidence block to support interpretation, not just to attach raw outputs."


def _main_report_summary_from_title(title: str) -> str:
    lower_title = _normalize_text(title or "").lower()
    for key, summary in MAIN_REPORT_SUMMARIES.items():
        if key.replace("_", " ") in lower_title or key in lower_title.replace(" ", "_"):
            return summary
    return "This evidence block summarises one of the main live outputs that supports the overall investment case."


def _main_report_use_from_title(title: str) -> str:
    lower_title = _normalize_text(title or "").lower()
    for key, summary in MAIN_REPORT_USE.items():
        if key.replace("_", " ") in lower_title or key in lower_title.replace(" ", "_"):
            return summary
    return "Use this evidence to explain what the live output is showing, why it matters, and whether it supports the core claim."


def _asset_name_label(asset_name: str) -> str:
    return _specific_asset_label(Path(asset_name).stem) or _humanize_slug(Path(asset_name).stem)


def _friendly_source_context_label(source_context: str) -> str:
    normalized = _normalize_text(source_context or "")
    if not normalized:
        return "the supporting evidence set"
    lowered = normalized.lower()
    if "part 1" in lowered:
        return "the parameter-stability evidence set"
    if "part 2" in lowered:
        return "the component-contribution evidence set"
    if "part 3" in lowered:
        return "the time-stability evidence set"
    if "part 4" in lowered:
        return "the simulation-based robustness evidence set"
    if "part 5" in lowered:
        return "the robustness summary set"
    if "main" in lowered and "report" in lowered:
        return "the main-program evidence set"
    return normalized


def _asset_source_line(asset_path: Path, source_context: str) -> str:
    suffix = asset_path.suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg"}:
        kind = "chart"
    elif suffix in {".csv", ".json"}:
        kind = "table"
    elif suffix == ".md":
        kind = "note"
    else:
        kind = "file"
    label = _asset_name_label(asset_path.name)
    context_label = _friendly_source_context_label(source_context)
    return f"Evidence source: {kind} '{label}' from {context_label}."


def _selected_evidence_line(asset_names: list[str]) -> str:
    clean_names = [_asset_name_label(str(name)) for name in asset_names if str(name).strip()]
    if not clean_names:
        return ""
    return f"Selected evidence: {', '.join(clean_names)}."


def _evidence_source_context_line(source_context: str) -> str:
    return f"Evidence source: {_friendly_source_context_label(source_context)}."


def _has_previous_update(payload: dict) -> bool:
    context_snapshot = payload.get("context_snapshot") or {}
    if not isinstance(context_snapshot, dict):
        return False
    previous_update = context_snapshot.get("previous_update") or {}
    if not isinstance(previous_update, dict):
        return False
    return bool(str(previous_update.get("report_id") or "").strip())


def _item_summary(item_label: str) -> str:
    return ITEM_SUMMARIES.get(
        item_label, f"This evidence block supports the written explanation for {item_label}."
    )


def _item_report_use(item_label: str) -> str:
    return ITEM_REPORT_USE.get(
        item_label,
        f"Use these outputs to explain what {item_label} shows, why it matters, and whether it strengthens or weakens the investment case.",
    )


def _latest_report_bundle(reports_dir: Path) -> Path | None:
    if not reports_dir.exists():
        return None
    candidates = [path for path in reports_dir.iterdir() if path.is_dir()]
    if not candidates:
        return None
    return sorted(candidates, key=lambda path: path.stat().st_mtime, reverse=True)[0]


def _preferred_mainline_report_bundle(base_dir: Path, reports_dir: Path) -> Path | None:
    search_roots = [
        base_dir / "outputs" / "robustness" / "sensitivity" / "reports",
        base_dir / "outputs" / "robustness" / "ablation" / "reports",
        reports_dir,
    ]
    candidates = [
        path
        for root in search_roots
        if root.exists()
        for path in root.iterdir()
        if path.is_dir() and (path / "report_summary.json").exists()
    ]
    mainline_candidates = [
        path for path in candidates if "cost_15bps_mainline" in path.name.lower()
    ]
    if mainline_candidates:
        return sorted(mainline_candidates, key=lambda path: path.stat().st_mtime, reverse=True)[0]
    return _latest_report_bundle(reports_dir)


def _iter_main_report_assets(report_dir: Path | None) -> list[Path]:
    if report_dir is None:
        return []
    assets = [
        path
        for path in report_dir.rglob("*")
        if path.is_file()
        and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".csv", ".md", ".json"}
    ]
    preferred_positions = {
        name.lower(): idx for idx, name in enumerate(PREFERRED_MAIN_REPORT_ORDER)
    }
    return sorted(
        assets,
        key=lambda path: (
            preferred_positions.get(path.name.lower(), 999),
            path.suffix.lower(),
            path.name.lower(),
        ),
    )


def _main_report_body_assets(report_dir: Path | None) -> list[Path]:
    appendix_names = {name.lower() for name in APPENDIX_MAIN_REPORT_ORDER}
    selected: list[Path] = []
    for asset in _iter_main_report_assets(report_dir):
        if asset.name.lower() in appendix_names:
            continue
        if asset.stem.lower() in MAIN_REPORT_CORE_EXCLUSIONS:
            continue
        selected.append(asset)
        if len(selected) >= MAX_MAIN_PROGRAM_ASSETS:
            break
    return selected


def _main_report_appendix_assets(report_dir: Path | None) -> list[Path]:
    appendix_names = {name.lower() for name in APPENDIX_MAIN_REPORT_ORDER}
    return [
        asset
        for asset in _iter_main_report_assets(report_dir)
        if asset.name.lower() in appendix_names
    ]


def _render_main_report_asset(document: Document, path: Path) -> None:
    stem = path.stem.lower()
    title = _humanize_slug(path.stem)
    explanation = next(
        (summary for key, summary in MAIN_REPORT_SUMMARIES.items() if key in stem),
        "This asset comes from the latest main-program report bundle and is included to ground the narrative in the current production outputs.",
    )
    document.add_heading(title, level=3)
    document.add_paragraph(explanation)
    suffix = path.suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg"}:
        document.add_picture(str(path), width=Inches(6.0))
    elif suffix == ".csv":
        _add_csv_table(document, path)
    elif suffix == ".json":
        _add_json_summary(document, path)
    elif suffix == ".md":
        _add_markdown_blocks(document, _safe_read_text(path))


def _selected_evidence_payload(payload: dict) -> dict:
    selected = payload.get("selected_evidence")
    return selected if isinstance(selected, dict) else {}


def _report_summary_payload(report_dir: Path | None) -> dict:
    if report_dir is None:
        return {}
    summary_path = report_dir / "report_summary.json"
    payload = _load_json(summary_path) if summary_path.exists() else None
    return payload if isinstance(payload, dict) else {}


def _summary_percent(payload: dict, key: str) -> str:
    value = payload.get(key)
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return ""
    return f"{float(value):.2f}%"


def _summary_number(payload: dict, key: str) -> str:
    value = payload.get(key)
    if not isinstance(value, (int, float)) or isinstance(value, bool):
        return ""
    return f"{float(value):.2f}"


def _fmt_percent_number(value: float | None) -> str:
    if value is None:
        return ""
    return f"{value:.2f}%"


def _fmt_ratio_number(value: float | None) -> str:
    if value is None:
        return ""
    return f"{value:.2f}"


def _row_float(row: dict[str, str], key: str) -> float | None:
    return _safe_float(str(row.get(key, "")))


def _read_csv_dicts(csv_path: Path) -> list[dict[str, str]]:
    if not csv_path.exists():
        return []
    try:
        with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
            return list(csv.DictReader(handle))
    except Exception:
        return []


def _base_dir_for_report_dir(report_dir: Path | None) -> Path | None:
    if report_dir is None:
        return None
    for parent in [report_dir, *report_dir.parents]:
        if (parent / "outputs").exists() and (parent / "scripts").exists():
            return parent
    return None


def _mainline_sensitivity_row(report_dir: Path | None) -> dict[str, str]:
    base_dir = _base_dir_for_report_dir(report_dir)
    if base_dir is None:
        return {}
    candidates = [
        base_dir
        / "outputs"
        / "robustness"
        / "sensitivity"
        / "summaries"
        / "deterministic_sensitivity_summary_completed_1_3_6.csv",
        base_dir
        / "outputs"
        / "robustness"
        / "sensitivity"
        / "summaries"
        / "deterministic_sensitivity_summary.csv",
    ]
    for csv_path in candidates:
        rows = _read_csv_dicts(csv_path)
        for row in rows:
            if str(row.get("scenario_key") or "").strip() == "cost_15bps_mainline":
                return row
    return {}


def _cost_sensitivity_rows(report_dir: Path | None) -> list[list[str]]:
    base_dir = _base_dir_for_report_dir(report_dir)
    if base_dir is None:
        return []
    csv_path = (
        base_dir
        / "outputs"
        / "robustness"
        / "report_evidence"
        / "part_1_deterministic"
        / "test_01_table.csv"
    )
    rows = _read_csv_dicts(csv_path)
    if not rows:
        return []
    table_rows = [["Cost assumption", "Annualized return", "Sharpe", "Implementation read-through"]]
    for row in rows:
        label = _friendly_scenario_value(
            "test_01", str(row.get("scenario_key") or row.get("scenario_label") or "")
        )
        annual_return = _fmt_percent_number(_row_float(row, "return.annualized_return"))
        sharpe = _fmt_ratio_number(_row_float(row, "risk_adjusted.sharpe_ratio"))
        cost_drag = _fmt_percent_number(_row_float(row, "portfolio.total_cost_drag"))
        static_excess = _fmt_percent_number(
            _row_float(row, "static_baseline.excess_return_annualized")
        )
        if not label:
            continue
        read_through_parts = []
        if cost_drag:
            read_through_parts.append(f"cost drag {cost_drag}")
        if static_excess:
            read_through_parts.append(f"excess vs static baseline {static_excess}")
        table_rows.append([label, annual_return, sharpe, "; ".join(read_through_parts)])
    return table_rows if len(table_rows) > 1 else []


def _benchmark_comparison_rows(report_dir: Path | None) -> list[list[str]]:
    row = _mainline_sensitivity_row(report_dir)
    if not row:
        return []
    strategy_total = _row_float(row, "return.total_return")
    strategy_ann = _row_float(row, "return.annualized_return")
    strategy_dd = _row_float(row, "risk.max_drawdown")
    strategy_sharpe = _row_float(row, "risk_adjusted.sharpe_ratio")
    table_rows = [["Series", "Annualized return", "Total return", "Risk / relative read-through"]]
    table_rows.append(
        [
            "Strategy",
            _fmt_percent_number(strategy_ann),
            _fmt_percent_number(strategy_total),
            f"Max drawdown {_fmt_percent_number(strategy_dd)}; Sharpe {_fmt_ratio_number(strategy_sharpe)}",
        ]
    )
    for series, label in [
        ("universe_ew", "Universe EW"),
        ("SPY", "SPY"),
        ("static_baseline", "Static baseline"),
    ]:
        excess_ann = _row_float(row, f"{series}.excess_return_annualized")
        excess_total = _row_float(row, f"{series}.excess_return_total")
        dd_delta = _row_float(row, f"{series}.max_drawdown_delta")
        tracking_error = _row_float(row, f"{series}.tracking_error")
        information_ratio = _row_float(row, f"{series}.information_ratio")
        benchmark_ann = (
            strategy_ann - excess_ann
            if strategy_ann is not None and excess_ann is not None
            else None
        )
        benchmark_total = (
            strategy_total - excess_total
            if strategy_total is not None and excess_total is not None
            else None
        )
        benchmark_dd = (
            strategy_dd - dd_delta if strategy_dd is not None and dd_delta is not None else None
        )
        read_through = []
        if benchmark_dd is not None:
            read_through.append(f"benchmark max drawdown {_fmt_percent_number(benchmark_dd)}")
        if information_ratio is not None:
            read_through.append(f"IR {_fmt_ratio_number(information_ratio)}")
        if tracking_error is not None:
            read_through.append(f"TE {_fmt_percent_number(tracking_error)}")
        table_rows.append(
            [
                label,
                _fmt_percent_number(benchmark_ann),
                _fmt_percent_number(benchmark_total),
                "; ".join(read_through),
            ]
        )
    return table_rows


def _add_benchmark_comparison_snapshot(document: Document, report_dir: Path | None) -> None:
    rows = _benchmark_comparison_rows(report_dir)
    if not rows:
        return
    document.add_heading("Strategy Vs Benchmarks Snapshot", level=4)
    document.add_paragraph(
        "This table turns the benchmark evidence into an investor-facing comparison. Benchmark rows are derived from the reported relative metrics, so they should be read as a compact comparison view rather than a separate backtest."
    )
    _add_flow_rows(document, rows)


def _add_cost_sensitivity_snapshot(document: Document, report_dir: Path | None) -> None:
    rows = _cost_sensitivity_rows(report_dir)
    if not rows:
        return
    document.add_heading("Gross-To-Net And Cost Sensitivity", level=4)
    document.add_paragraph(
        "This table connects turnover to investability. It shows whether the strategy still has room to absorb higher trading-cost assumptions, instead of treating turnover as a decorative implementation chart."
    )
    _add_flow_rows(document, rows)


def _asset_key_from_block(title: str, asset_names: list[str]) -> str:
    candidates = [_normalize_text(title or "").lower().replace(" ", "_")]
    candidates.extend(Path(name).stem.lower() for name in asset_names)
    return " ".join(candidates)


def _is_report_summary_block(block: dict[str, object]) -> bool:
    title = _normalize_text(str(block.get("title") or ""))
    asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
    return "report_summary" in _asset_key_from_block(title, asset_names)


def _report_summary_block(selected_evidence: dict[str, object]) -> dict[str, object]:
    main_blocks = (
        selected_evidence.get("main_report") if isinstance(selected_evidence, dict) else None
    )
    if isinstance(main_blocks, list):
        for block in main_blocks:
            if isinstance(block, dict) and _is_report_summary_block(block):
                return block
    return {}


def _main_report_override_analysis(
    title: str, asset_names: list[str], report_dir: Path | None
) -> str | None:
    block_key = _asset_key_from_block(title, asset_names)
    summary = _report_summary_payload(report_dir)
    if "report_summary" in block_key:
        annual_return = _summary_percent(summary, "annualized_return")
        excess_return = _summary_percent(summary, "excess_return_vs_primary")
        drawdown = _summary_percent(summary, "max_drawdown")
        volatility = _summary_percent(summary, "annualized_volatility")
        sharpe = _summary_number(summary, "sharpe_ratio")
        benchmark = _friendly_benchmark_label(summary.get("primary_benchmark") or "SPY")
        turnover = _summary_percent(summary, "avg_monthly_turnover")
        parts = ["The mainline summary is the authoritative headline source for this report."]
        if annual_return and excess_return:
            parts.append(
                f"It reports {annual_return} annualized return and {excess_return} annualized excess return versus {benchmark}, so the backtest evidence supports a positive active return claim."
            )
        if sharpe and drawdown and volatility:
            parts.append(
                f"Risk is still equity-like: Sharpe is {sharpe}, maximum drawdown is {drawdown}, and annualized volatility is {volatility}."
            )
        if turnover:
            parts.append(
                f"Average monthly turnover is {turnover}, so implementation activity should be treated as a material monitoring point rather than a minor footnote."
            )
        return " ".join(parts)
    if "turnover_and_cost" in block_key:
        monthly_turnover = _summary_percent(summary, "avg_monthly_turnover")
        annual_turnover = _summary_percent(summary, "annualized_turnover_ratio")
        cost_bps = summary.get("transaction_cost_bps")
        cost_text = (
            f"{float(cost_bps):.0f} bps" if isinstance(cost_bps, (int, float)) else "the stated"
        )
        excess_return = _summary_percent(summary, "excess_return_vs_primary")
        return (
            f"The turnover evidence is important because implementability is the main practical caveat. "
            f"The mainline run uses a {cost_text} transaction-cost assumption"
            f"{f' with average monthly turnover of {monthly_turnover}' if monthly_turnover else ''}"
            f"{f' and annualized turnover of {annual_turnover}' if annual_turnover else ''}. "
            "That level of trading activity is material, not negligible. "
            f"The strategy still outperforms after the cost assumption{f', but the {excess_return} annualized excess return should be monitored against higher-cost or lower-liquidity scenarios' if excess_return else ', but the excess return should be monitored against higher-cost or lower-liquidity scenarios'}."
        )
    if "regime_return_summary" in block_key:
        return (
            "The regime view should be read as a conditional-performance check, not as proof that the strategy is defensive. "
            "Normal-market periods support the return case, while stress periods can still produce negative absolute returns even when benchmark-relative excess return remains positive. "
            "This supports regime-aware monitoring and keeps the risk conclusion appropriately cautious."
        )
    if "latest_sector_risk_contribution" in block_key:
        return (
            "This exhibit is a sector risk-contribution view, not a factor-allocation chart. "
            "Use it to assess whether portfolio risk is overly concentrated in one sector; factor allocation claims should be supported by separate factor exposure or attribution evidence. "
            "Its role in the report is therefore to support concentration-risk monitoring rather than to validate the quality, value, market-technical, or dividend factor mix directly."
        )
    return None


def _turnover_monitoring_sentence(summary: dict) -> str:
    monthly_turnover = _summary_percent(summary, "avg_monthly_turnover")
    annual_turnover = _summary_percent(summary, "annualized_turnover_ratio")
    gross_return = _summary_percent(summary, "gross_annualized_return")
    net_return = _summary_percent(summary, "annualized_return")
    cost_bps = summary.get("transaction_cost_bps")
    cost_text = (
        f"{float(cost_bps):.0f} bps"
        if isinstance(cost_bps, (int, float))
        else "the stated transaction-cost"
    )
    gross_to_net = ""
    if gross_return and net_return:
        gross_to_net = f" Gross annualized return is {gross_return} before costs and {net_return} after the cost assumption, so the report should treat the gross-to-net gap as part of the investment case."
    if monthly_turnover and annual_turnover:
        return (
            f"Implementation cost is included through a {cost_text} assumption, but average monthly turnover of {monthly_turnover} "
            f"and annualized turnover of {annual_turnover} make trading intensity a material monitoring point."
            f"{gross_to_net}"
        )
    if monthly_turnover:
        return (
            f"Implementation cost is included through a {cost_text} assumption, but average monthly turnover of {monthly_turnover} "
            "makes trading intensity a material monitoring point."
            f"{gross_to_net}"
        )
    return ""


def _align_section_with_mainline_metrics(title: str, body: str, summary: dict) -> str:
    text = _normalize_text(body or "")
    if not text or not summary:
        return text
    percent_value = r"-?\d+(?:\.\d+)?%"
    decimal_value = r"-?\d+(?:\.\d+)?"

    def replace_metric(target: str, pattern: str, value: str) -> str:
        if not value:
            return target
        return re.sub(
            pattern,
            lambda match: f"{match.group(1)}{value}{match.group(2) if match.lastindex and match.lastindex > 1 else ''}",
            target,
            flags=re.IGNORECASE,
        )

    def align_headline_paragraph(paragraph: str) -> str:
        annualized = _summary_percent(summary, "annualized_return")
        total = _summary_percent(summary, "total_return")
        excess = _summary_percent(summary, "excess_return_vs_primary")
        drawdown = _summary_percent(summary, "max_drawdown")
        volatility = _summary_percent(summary, "annualized_volatility")
        sharpe = _format_json_value("sharpe_ratio", summary.get("sharpe_ratio"))
        paragraph = replace_metric(
            paragraph,
            rf"(\bannuali[sz]ed return(?:\s+(?:of|around|near|is|was|at|exceeding))?\s*){percent_value}",
            annualized,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\bcumulative return(?:\s+(?:of|around|near|is|was|at))?\s*){percent_value}",
            total,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\b(?:excess|benchmark-relative excess) return(?:\s+(?:of|around|near|is|was|at))?\s*){percent_value}(\s+(?:above|versus|vs|relative to)[^.\n]{{0,100}}(?:primary benchmark|Universe EW|equal-weighted universe|benchmark))",
            excess,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\boutperformed[^.\n]{{0,100}}(?:Universe EW|primary benchmark|equal-weighted universe)[^.\n]{{0,40}}\bby\s*){percent_value}(\s+annuali[sz]ed)?",
            excess,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\bmaximum drawdown(?:\s+(?:of|around|near|is|was|at))?\s*){percent_value}",
            drawdown,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\bvolatility(?:\s+(?:of|around|near|is|was|at))?\s*){percent_value}",
            volatility,
        )
        paragraph = replace_metric(
            paragraph,
            rf"(\bSharpe ratio(?:\s+(?:of approximately|of|around|near|is|was|at))?\s*){decimal_value}\b",
            sharpe,
        )
        return paragraph

    if title in {"Executive Summary", "Backtest Results", "Limitations And Monitoring Signals"}:
        parts = re.split(r"(\n\s*\n)", text, maxsplit=1)
        parts[0] = align_headline_paragraph(parts[0])
        text = "".join(parts)
    excess_return = _summary_percent(summary, "excess_return_vs_primary")
    if excess_return and title in {"Executive Summary", "Backtest Results"}:
        text = replace_metric(
            text,
            rf"(\bbenchmark-relative excess return(?:\s+(?:of|around|near|is|was|at))?\s*){percent_value}(\s+versus\s+Universe EW)",
            excess_return,
        )
    turnover_sentence = _turnover_monitoring_sentence(summary)
    if turnover_sentence and title == "Backtest Results":
        text = re.sub(
            r"Turnover and transaction cost impacts are implicitly accounted for in robustness tests, with cost perturbation scenarios showing only modest degradation in performance metrics\.",
            turnover_sentence,
            text,
        )
        if "turnover" not in text.lower():
            text = f"{text}\n\n{turnover_sentence}"
    if turnover_sentence and title == "Limitations And Monitoring Signals":
        text = re.sub(
            r"Turnover and transaction costs, while tested, remain a potential drag under adverse market conditions\.",
            turnover_sentence,
            text,
        )
        if "turnover" not in text.lower():
            text = f"{text}\n\n{turnover_sentence}"
    if turnover_sentence and title == "Executive Summary" and "turnover" not in text.lower():
        text = f"{text}\n\n{turnover_sentence}"
    return text


def _render_asset_content(document: Document, path: Path) -> None:
    suffix = path.suffix.lower()
    if suffix in {".png", ".jpg", ".jpeg"}:
        document.add_picture(str(path), width=Inches(6.0))
        caption = _brief_image_caption(path.name)
        if caption:
            document.add_paragraph(caption)
    elif suffix == ".csv":
        if _csv_has_meaningful_content(path):
            _add_csv_table(document, path, compact=True)
            caption = _brief_table_caption(path.name)
            if caption:
                document.add_paragraph(caption)
    elif suffix == ".json":
        _add_json_summary(document, path)
    elif suffix == ".md":
        note_text = _clean_note_text(_safe_read_text(path))
        if note_text:
            _add_markdown_blocks(document, note_text)


def _render_main_report_appendix_summary(document: Document, report_dir: Path | None) -> bool:
    if report_dir is None:
        return False
    summary_path = report_dir / "report_summary.json"
    if not summary_path.exists():
        return False
    document.add_heading("Report Summary", level=3)
    document.add_paragraph(
        "This summary gives a compact reference point for the latest main-program reporting run without repeating the main discussion."
    )
    lines = _json_summary_lines(summary_path)
    if not lines:
        document.add_paragraph(
            "A concise structured summary is unavailable for this reporting run."
        )
        return True
    for line in lines[:14]:
        document.add_paragraph(line, style="List Bullet")
    return True


def _render_opening_report_summary(
    document: Document,
    report_dir: Path | None,
    selected_evidence: dict[str, object],
    evidence_dir: Path | None = None,
) -> bool:
    if report_dir is None:
        return False
    summary_path = report_dir / "report_summary.json"
    if not summary_path.exists():
        return False
    document.add_heading("Report Summary", level=1)
    document.add_paragraph(
        "This opening snapshot anchors the report in the mainline backtest run before the narrative sections interpret the strategy, results, risk, and robustness evidence."
    )
    _add_opening_metric_snapshot(document, summary_path, evidence_dir)
    takeaway_para = document.add_paragraph(_opening_takeaway(summary_path))
    takeaway_para.paragraph_format.space_before = Pt(5)
    return True


def _normalized_heading_text(text: str) -> str:
    return re.sub(r"[^a-z0-9]+", " ", (text or "").lower()).strip()


def _asset_caption_text(asset_path: Path) -> str:
    suffix = asset_path.suffix.lower()
    label = _specific_asset_label(asset_path.stem) or _humanize_slug(asset_path.stem)
    if suffix in {".png", ".jpg", ".jpeg"}:
        return f"Fig. {label}."
    if suffix in {".csv", ".json"}:
        return f"Tbl. {label}."
    if suffix == ".md":
        return f"Note. {label}."
    return label


def _brief_image_caption(filename: str) -> str | None:
    lower_name = filename.lower()
    stem = Path(filename).stem.lower()
    for key, caption in SPECIFIC_IMAGE_CAPTIONS.items():
        if key in stem:
            return caption
    for keyword, caption in BRIEF_IMAGE_CAPTIONS.items():
        if keyword in lower_name:
            return caption
    return None


def _brief_table_caption(filename: str) -> str | None:
    lower_name = filename.lower()
    stem = Path(filename).stem.lower()
    for key, caption in SPECIFIC_TABLE_CAPTIONS.items():
        if key in stem:
            return caption
    if "table" in lower_name:
        return "This table summarises the selected evidence in a compact form so the reader can compare scenarios and outcomes directly."
    return None


def _maybe_add_asset_heading(document: Document, asset_path: Path, parent_title: str) -> None:
    asset_title = _humanize_slug(asset_path.stem)
    if _normalized_heading_text(asset_title) != _normalized_heading_text(parent_title):
        document.add_paragraph(_asset_caption_text(asset_path))


def _body_assets_from_names(asset_names: list[str], base_dir: Path) -> list[Path]:
    assets: list[Path] = []
    for asset_name in asset_names:
        candidate = base_dir / asset_name
        if candidate.exists():
            assets.append(candidate)
    figures = [asset for asset in assets if asset.suffix.lower() in {".png", ".jpg", ".jpeg"}]
    tables = [asset for asset in assets if asset.suffix.lower() in {".csv", ".json"}]
    notes = [asset for asset in assets if asset.suffix.lower() == ".md"]
    selected: list[Path] = []
    if figures:
        selected.append(figures[0])
    if tables:
        selected.append(tables[0])
    if not selected and notes:
        selected.append(notes[0])
    return selected


def _appendix_assets_from_names(asset_names: list[str], base_dir: Path) -> list[Path]:
    assets: list[Path] = []
    for asset_name in asset_names:
        candidate = base_dir / asset_name
        if candidate.exists():
            assets.append(candidate)
    body_assets = set(_body_assets_from_names(asset_names, base_dir))
    appendix_assets = [asset for asset in assets if asset not in body_assets]
    if appendix_assets:
        return appendix_assets
    return []


def _selected_block_has_appendix_assets(block: dict[str, object], base_dir: Path) -> bool:
    asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
    return bool(_appendix_assets_from_names(asset_names, base_dir))


def _render_selected_main_report_block(
    document: Document,
    block: dict[str, object],
    report_dir: Path | None,
    *,
    heading_level: int = 3,
    display_title: str | None = None,
    include_heading: bool = True,
) -> None:
    raw_title = _normalize_text(str(block.get("title") or "Main report evidence"))
    title = _normalize_text(display_title or raw_title)
    asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
    block_key = _asset_key_from_block(raw_title, asset_names)
    raw_analysis, inline_interpretation = _split_evidence_analysis_text(
        str(block.get("analysis") or "")
    )
    analysis = _main_report_override_analysis(
        raw_title, asset_names, report_dir
    ) or _sanitize_evidence_analysis(raw_analysis)
    investor_interpretation = _block_investor_interpretation(
        block,
        title,
        asset_names,
        inline_interpretation=inline_interpretation,
        report_dir=report_dir,
    )
    if any(Path(name).stem.lower() in MAIN_REPORT_CORE_EXCLUSIONS for name in asset_names):
        return
    if include_heading:
        document.add_heading(title, level=heading_level)
    plain_note = (
        _plain_language_evidence_note(raw_title)
        or _plain_language_evidence_note(title)
        or _main_report_summary_from_title(raw_title)
    )
    if plain_note:
        document.add_paragraph(plain_note)
    if analysis:
        document.add_paragraph(analysis)
    _add_investor_interpretation(document, investor_interpretation)
    if report_dir is None:
        return
    for asset_path in _body_assets_from_names(asset_names, report_dir):
        _maybe_add_asset_heading(document, asset_path, title)
        _render_asset_content(document, asset_path)
    if "report_summary" in block_key:
        _add_benchmark_comparison_snapshot(document, report_dir)
    elif "turnover_and_cost" in block_key:
        _add_cost_sensitivity_snapshot(document, report_dir)


def _render_selected_part_body(
    document: Document, part_label: str, part_dir: Path, selected_blocks: list[dict[str, object]]
) -> None:
    pretty_part = PART_DIR_LABELS.get(part_dir.name, part_label)
    document.add_heading(pretty_part, level=2)
    document.add_paragraph(_part_summary(part_label))
    document.add_paragraph(_part_report_use(part_label))
    for block in selected_blocks:
        raw_item = _normalize_text(str(block.get("item") or block.get("title") or "Evidence item"))
        title = _reader_friendly_item_label(raw_item)
        raw_analysis, inline_interpretation = _split_evidence_analysis_text(
            str(block.get("analysis") or "")
        )
        analysis = _sanitize_evidence_analysis(raw_analysis)
        item_label = raw_item
        asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
        investor_interpretation = _block_investor_interpretation(
            block,
            title,
            asset_names,
            inline_interpretation=inline_interpretation,
            part_label=part_label,
        )
        document.add_heading(title, level=3)
        plain_note = (
            _plain_language_evidence_note(raw_item)
            or _plain_language_evidence_note(title)
            or _normalize_text(_item_summary(item_label))
        )
        if plain_note:
            document.add_paragraph(plain_note)
        selected_line = _selected_evidence_line(asset_names)
        if selected_line:
            document.add_paragraph(selected_line)
        document.add_paragraph(_evidence_source_context_line(pretty_part))
        if "test 2" in raw_item.lower() or any("test_02" in name.lower() for name in asset_names):
            document.add_paragraph(
                "Window-test base case is a scenario-specific robustness rerun and may differ slightly from the headline production mainline."
            )
        if analysis:
            document.add_paragraph(analysis)
        _add_investor_interpretation(document, investor_interpretation)
        for asset_path in _body_assets_from_names(asset_names, part_dir):
            _maybe_add_asset_heading(document, asset_path, title)
            _render_asset_content(document, asset_path)


def _render_selected_part_appendix(
    document: Document, part_label: str, part_dir: Path, selected_blocks: list[dict[str, object]]
) -> bool:
    pretty_part = PART_DIR_LABELS.get(part_dir.name, part_label)
    rendered_blocks: list[tuple[str, list[Path]]] = []
    for block in selected_blocks:
        title = str(block.get("title") or block.get("item") or "Evidence item")
        asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
        appendix_assets = _appendix_assets_from_names(asset_names, part_dir)
        if not appendix_assets:
            continue
        rendered_blocks.append((title, appendix_assets))
    if not rendered_blocks:
        return False
    document.add_heading(pretty_part, level=2)
    for title, appendix_assets in rendered_blocks:
        document.add_heading(title, level=3)
        for asset_path in appendix_assets:
            _maybe_add_asset_heading(document, asset_path, title)
            _render_asset_content(document, asset_path)
    return True


def _read_report_evidence_index(evidence_dir: Path) -> list[dict[str, str]]:
    index_path = evidence_dir / "REPORT_EVIDENCE_INDEX.csv"
    if not index_path.exists():
        return []
    with index_path.open("r", encoding="utf-8-sig", newline="") as handle:
        return list(csv.DictReader(handle))


def _normalize_part_dir(part_label: str) -> str:
    mapping = {
        "Part 1": "part_1_deterministic",
        "Part 2": "part_2_ablation",
        "Part 3": "part_3_subperiod",
        "Part 4": "part_4_stochastic",
        "Part 5": "part_5_dashboard_and_conclusions",
    }
    return mapping.get(part_label.strip(), "")


def _prefix_candidates(*paths: Path) -> list[str]:
    prefixes: list[str] = []
    for path in paths:
        stem = path.stem.lower()
        prefixes.append(stem)
        for suffix in (
            "_notes",
            "_table",
            "_chart",
            "_summary",
            "_nav_reference",
            "_bootstrap_sharpe_hist",
            "_cost_sigma30_sharpe_hist",
            "_factor_neighbourhood_summary",
            "_oos_excess_return_hist",
            "_parametric_sharpe_hist",
            "_sample_paths",
            "_report_ready_summary",
            "_report_ready_notes",
        ):
            if stem.endswith(suffix):
                prefixes.append(stem[: -len(suffix)])
    return [prefix for prefix in dict.fromkeys(prefixes) if prefix]


def _collect_related_assets(part_dir: Path, row: dict[str, str]) -> list[Path]:
    named_assets = []
    for key in ("table", "figure", "notes"):
        raw_name = (row.get(key) or "").strip()
        if raw_name:
            candidate = part_dir / raw_name
            if candidate.exists():
                named_assets.append(candidate)
    prefixes = _prefix_candidates(*named_assets)
    related: list[Path] = []
    seen: set[Path] = set()
    for asset in named_assets:
        if asset not in seen:
            related.append(asset)
            seen.add(asset)
    for path in sorted(part_dir.iterdir(), key=lambda item: item.name.lower()):
        if not path.is_file():
            continue
        lower_name = path.name.lower()
        if lower_name in {
            "manifest.json",
            "report_evidence_index.csv",
            "report_evidence_index.md",
            "robustness_report_evidence_pack.md",
        }:
            continue
        stem = path.stem.lower()
        if any(stem.startswith(prefix) for prefix in prefixes) and path not in seen:
            related.append(path)
            seen.add(path)
    return related


def _safe_float(value: str) -> float | None:
    text = (value or "").strip().replace(",", "")
    if not text:
        return None
    if text.endswith("%"):
        text = text[:-1]
    try:
        return float(text)
    except ValueError:
        return None


def _csv_numeric_profile(csv_path: Path, max_rows: int = 80) -> dict[str, float]:
    numeric_values: list[float] = []
    row_count = 0
    try:
        with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
            reader = csv.reader(handle)
            for idx, row in enumerate(reader):
                if idx >= max_rows:
                    break
                row_count += 1
                for cell in row:
                    parsed = _safe_float(cell)
                    if parsed is not None:
                        numeric_values.append(parsed)
    except Exception:
        return {
            "row_count": 0.0,
            "numeric_count": 0.0,
            "distinct_numeric_count": 0.0,
            "max_abs": 0.0,
            "spread": 0.0,
        }
    if not numeric_values:
        return {
            "row_count": float(row_count),
            "numeric_count": 0.0,
            "distinct_numeric_count": 0.0,
            "max_abs": 0.0,
            "spread": 0.0,
        }
    return {
        "row_count": float(row_count),
        "numeric_count": float(len(numeric_values)),
        "distinct_numeric_count": float(len({round(value, 6) for value in numeric_values})),
        "max_abs": max(abs(value) for value in numeric_values),
        "spread": max(numeric_values) - min(numeric_values),
    }


def _score_evidence_row(part_dir: Path, row: dict[str, str]) -> float:
    assets = _collect_related_assets(part_dir, row)
    score = 0.0
    csv_profiles: list[dict[str, float]] = []
    for path in assets:
        suffix = path.suffix.lower()
        stem = path.stem.lower()
        if suffix == ".md":
            score += 2.5
        elif suffix in {".png", ".jpg", ".jpeg"}:
            score += 3.0
            if "nav_reference" in stem:
                score += 0.5
        elif suffix == ".csv":
            profile = _csv_numeric_profile(path)
            csv_profiles.append(profile)
            score += 1.5
            score += min(3.0, profile["numeric_count"] / 12.0)
            score += min(2.5, profile["distinct_numeric_count"] / 8.0)
            if profile["spread"] > 0:
                score += min(2.5, math.log10(profile["spread"] + 1.0))
            if profile["max_abs"] > 0:
                score += min(2.0, math.log10(profile["max_abs"] + 1.0))
        elif suffix == ".json":
            score += 1.0
    item_label = (row.get("item") or "").strip()
    if item_label in {"Regime decomposition", "Fixed windows", "Dashboard"}:
        score += 1.0
    if item_label.lower().startswith("test"):
        score += 0.5
    return score


def _pick_assets_for_report(asset_paths: list[Path]) -> list[Path]:
    if not asset_paths:
        return []
    notes: list[Path] = []
    figures: list[Path] = []
    nav_refs: list[Path] = []
    tables: list[Path] = []
    other: list[Path] = []
    for path in asset_paths:
        stem = path.stem.lower()
        suffix = path.suffix.lower()
        if suffix == ".md":
            notes.append(path)
        elif "nav_reference" in stem:
            nav_refs.append(path)
        elif suffix in {".png", ".jpg", ".jpeg"}:
            figures.append(path)
        elif suffix == ".csv":
            tables.append(path)
        else:
            other.append(path)
    selected: list[Path] = []
    if notes:
        selected.append(notes[0])
    if figures:
        selected.append(figures[0])
    elif nav_refs:
        selected.append(nav_refs[0])
    if tables:
        selected.append(tables[0])
    for path in other:
        if path not in selected:
            selected.append(path)
    return selected


def _selected_rows_for_part(part_label: str, rows: list[dict[str, str]]) -> list[dict[str, str]]:
    selection_limit = PART_SELECTION_LIMITS.get(part_label)
    if selection_limit is None:
        return rows
    if selection_limit <= 0:
        return []
    if len(rows) <= selection_limit:
        return rows
    scored_rows = []
    for idx, row in enumerate(rows):
        score = _score_evidence_row(Path(row.get("__part_dir__", "")), row)
        item_label = (row.get("item") or "").strip().lower()
        scored_rows.append((score, item_label, idx, row))
    scored_rows.sort(key=lambda item: (-item[0], item[1], item[2]))
    return [row for _, _, _, row in scored_rows[:selection_limit]]


def _render_evidence_asset(document: Document, path: Path) -> None:
    suffix = path.suffix.lower()
    label = _specific_asset_label(path.stem) or _humanize_slug(path.stem)
    document.add_heading(label, level=4)
    if suffix in {".png", ".jpg", ".jpeg"}:
        document.add_picture(str(path), width=Inches(6.0))
    elif suffix == ".csv":
        if not _csv_has_meaningful_content(path):
            return
        _add_csv_table(document, path, compact=True)
    elif suffix == ".json":
        _add_json_summary(document, path)
    elif suffix == ".md":
        note_text = _clean_note_text(_safe_read_text(path))
        _add_markdown_blocks(document, note_text)


def _render_part_body(
    document: Document, part_label: str, part_dir: Path, part_rows: list[dict[str, str]]
) -> None:
    pretty_part = PART_DIR_LABELS.get(part_dir.name, part_label)
    document.add_heading(pretty_part, level=2)
    document.add_paragraph(_part_summary(part_label))
    document.add_paragraph(_part_report_use(part_label))
    for row in part_rows:
        item_label = (row.get("item") or "").strip() or "Evidence item"
        display_label = _reader_friendly_item_label(item_label)
        document.add_heading(display_label, level=3)
        plain_note = _plain_language_evidence_note(
            _humanize_slug(item_label)
        ) or _plain_language_evidence_note(item_label)
        if plain_note:
            document.add_paragraph(plain_note)
        document.add_paragraph(_normalize_text(_item_summary(item_label)))
        document.add_paragraph(_normalize_text(_item_report_use(item_label)))
        if not part_dir.exists():
            continue
        for asset in _pick_assets_for_report(_collect_related_assets(part_dir, row)):
            _render_evidence_asset(document, asset)


def _render_part_appendix(
    document: Document, part_label: str, part_dir: Path, part_rows: list[dict[str, str]]
) -> bool:
    pretty_part = PART_DIR_LABELS.get(part_dir.name, part_label)
    rendered_rows: list[tuple[str, list[Path]]] = []
    for row in part_rows:
        item_label = (row.get("item") or "").strip() or "Evidence item"
        if not part_dir.exists():
            continue
        body_assets = set(_pick_assets_for_report(_collect_related_assets(part_dir, row)))
        appendix_assets = [
            asset for asset in _collect_related_assets(part_dir, row) if asset not in body_assets
        ]
        if not appendix_assets:
            continue
        rendered_rows.append((_reader_friendly_item_label(item_label), appendix_assets))
    if not rendered_rows:
        return False
    document.add_heading(pretty_part, level=2)
    for item_label, appendix_assets in rendered_rows:
        document.add_heading(item_label, level=3)
        for asset in appendix_assets:
            _render_evidence_asset(document, asset)
    return True


def _part_rows_have_appendix_assets(part_dir: Path, part_rows: list[dict[str, str]]) -> bool:
    for row in part_rows:
        body_assets = set(_pick_assets_for_report(_collect_related_assets(part_dir, row)))
        appendix_assets = [
            asset for asset in _collect_related_assets(part_dir, row) if asset not in body_assets
        ]
        if appendix_assets:
            return True
    return False


MAIN_RESULT_BLOCK_ORDER = [
    "nav_vs_benchmarks",
    "drawdown_comparison",
    "turnover_and_cost",
    "regime_return_summary",
    "latest_sector_risk_contribution",
]


MAIN_RESULT_SECTION_TITLES = {
    "nav_vs_benchmarks": "Return And Benchmark Performance",
    "drawdown_comparison": "Downside Risk And Drawdown",
    "turnover_and_cost": "Implementation Cost And Turnover",
    "regime_return_summary": "Market Regime Performance",
    "latest_sector_risk_contribution": "Exposure And Sector Risk",
}


def _prepared_section_body(
    title: str, sections: dict, payload: dict, mainline_summary: dict
) -> str:
    body = _sanitize_section_body(title, (sections.get(title) or "").strip())
    if _section_needs_fallback(title, body):
        body = _fallback_section_body(title, payload)
    return _align_section_with_mainline_metrics(title, body, mainline_summary)


def _render_narrative_section(
    document: Document,
    title: str,
    body: str,
    *,
    level: int,
    display_title: str | None = None,
    include_intro: bool = True,
) -> bool:
    if not body:
        return False
    document.add_heading(display_title or SECTION_DISPLAY_TITLES.get(title, title), level=level)
    intro = SECTION_INTROS.get(title) if include_intro else None
    if intro:
        document.add_paragraph(intro)
    _add_markdown_blocks(document, body)
    return True


def _render_narrative_group(
    document: Document,
    group_title: str,
    group_intro: str,
    section_titles: list[str],
    sections: dict,
    payload: dict,
    mainline_summary: dict,
) -> bool:
    prepared_sections = [
        (title, _prepared_section_body(title, sections, payload, mainline_summary))
        for title in section_titles
    ]
    prepared_sections = [(title, body) for title, body in prepared_sections if body]
    if not prepared_sections:
        return False
    document.add_heading(group_title, level=1)
    if group_intro:
        document.add_paragraph(group_intro)
    for title, body in prepared_sections:
        _render_narrative_section(document, title, body, level=2)
    return True


def _selected_main_report_blocks(selected_evidence: dict[str, object]) -> list[dict[str, object]]:
    main_blocks = (
        selected_evidence.get("main_report") if isinstance(selected_evidence, dict) else None
    )
    if not isinstance(main_blocks, list):
        return []
    return [
        block
        for block in main_blocks
        if isinstance(block, dict) and not _is_report_summary_block(block)
    ]


def _main_result_key(block: dict[str, object]) -> str:
    title = _normalize_text(str(block.get("title") or ""))
    asset_names = [str(name) for name in (block.get("asset_names") or []) if str(name).strip()]
    block_key = _asset_key_from_block(title, asset_names)
    for key in MAIN_RESULT_BLOCK_ORDER:
        if key in block_key:
            return key
    return ""


def _ordered_main_result_blocks(blocks: list[dict[str, object]]) -> list[dict[str, object]]:
    order = {key: idx for idx, key in enumerate(MAIN_RESULT_BLOCK_ORDER)}
    return sorted(
        blocks,
        key=lambda block: (
            order.get(_main_result_key(block), len(order)),
            str(block.get("title") or ""),
        ),
    )


def _render_backtest_results_analysis(
    document: Document,
    sections: dict,
    payload: dict,
    mainline_summary: dict,
    report_dir: Path | None,
    selected_evidence: dict[str, object],
) -> bool:
    result_body = _prepared_section_body("Backtest Results", sections, payload, mainline_summary)
    risk_body = _prepared_section_body(
        "Risk, Regime And Exposure Analysis", sections, payload, mainline_summary
    )
    selected_blocks = _ordered_main_result_blocks(_selected_main_report_blocks(selected_evidence))
    fallback_assets = (
        [
            asset
            for asset in _main_report_body_assets(report_dir)
            if asset.name.lower() != "report_summary.json"
        ]
        if report_dir is not None
        else []
    )
    if not any([result_body, risk_body, selected_blocks, fallback_assets]):
        return False

    document.add_heading("Backtest Results And Analysis", level=1)
    document.add_paragraph(
        "This section treats results, supporting exhibits, and interpretation as one analysis flow. Each subsection states what the result means and then shows the relevant output, so the reader does not have to jump between a result section and a separate evidence section."
    )
    if result_body:
        _render_narrative_section(
            document,
            "Backtest Results",
            result_body,
            level=2,
            display_title="Headline Performance Interpretation",
        )

    risk_body_rendered = False
    if selected_blocks:
        for block in selected_blocks:
            result_key = _main_result_key(block)
            if (
                risk_body
                and not risk_body_rendered
                and result_key in {"regime_return_summary", "latest_sector_risk_contribution"}
            ):
                _render_narrative_section(
                    document,
                    "Risk, Regime And Exposure Analysis",
                    risk_body,
                    level=2,
                    display_title="Regime And Exposure Interpretation",
                )
                risk_body_rendered = True
            display_title = MAIN_RESULT_SECTION_TITLES.get(result_key) or str(
                block.get("title") or "Backtest Output"
            )
            document.add_heading(display_title, level=2)
            _render_selected_main_report_block(
                document,
                block,
                report_dir,
                display_title=display_title,
                include_heading=False,
            )
            if result_key == "nav_vs_benchmarks":
                _add_benchmark_comparison_snapshot(document, report_dir)
    elif fallback_assets:
        for asset in fallback_assets:
            _render_main_report_asset(document, asset)

    if risk_body and not risk_body_rendered:
        _render_narrative_section(
            document,
            "Risk, Regime And Exposure Analysis",
            risk_body,
            level=2,
            display_title="Regime And Exposure Interpretation",
        )
    return True


def _concise_limitations_body(body: str) -> str:
    blocks = [
        block.strip()
        for block in re.split(r"\n\s*\n", _normalize_text(body or ""))
        if block.strip()
    ]
    if len(blocks) <= 3:
        return "\n\n".join(blocks)
    return "\n\n".join(blocks[:3])


def build_docx(payload: dict, output_docx: Path) -> None:
    base_dir = Path(str(payload.get("base_dir") or output_docx.parent))
    reports_dir = base_dir / "outputs" / "reports"
    evidence_dir = base_dir / "outputs" / "robustness" / "report_evidence"

    document = Document()
    _configure_document_styles(document)
    document.add_heading("Investment Portfolio Analysis Report", level=0)
    report_id = str(payload.get("report_id", "n/a"))
    generated_at = str(payload.get("generated_at", "n/a"))
    model_name = str(payload.get("model", "n/a"))
    document.add_paragraph(
        f"Generated on {generated_at}. "
        f"Update reference: {report_id}. "
        f"Model used: {model_name}."
    )

    sections = payload.get("sections") or {}
    selected_evidence = _selected_evidence_payload(payload)
    latest_bundle = _preferred_mainline_report_bundle(base_dir, reports_dir)
    mainline_summary = _report_summary_payload(latest_bundle)
    _render_opening_report_summary(document, latest_bundle, selected_evidence, evidence_dir)

    rendered_any_section = False
    executive_body = _prepared_section_body(
        "Executive Summary", sections, payload, mainline_summary
    )
    rendered_any_section = (
        _render_narrative_section(document, "Executive Summary", executive_body, level=1)
        or rendered_any_section
    )
    rendered_any_section = (
        _render_narrative_group(
            document,
            "Strategy And Backtest Design",
            "This section separates the investment process and testing setup from the performance discussion, so the reader can first understand what was built and how it was evaluated.",
            ["Strategy And Portfolio Construction", "Backtest Design"],
            sections,
            payload,
            mainline_summary,
        )
        or rendered_any_section
    )
    rendered_any_section = (
        _render_backtest_results_analysis(
            document,
            sections,
            payload,
            mainline_summary,
            latest_bundle,
            selected_evidence,
        )
        or rendered_any_section
    )

    limitations_body = _concise_limitations_body(
        _prepared_section_body(
            "Limitations And Monitoring Signals", sections, payload, mainline_summary
        )
    )
    rendered_any_section = (
        _render_narrative_section(
            document,
            "Limitations And Monitoring Signals",
            limitations_body,
            level=1,
        )
        or rendered_any_section
    )

    evidence_rows = _read_report_evidence_index(evidence_dir)
    appendix_rows_by_part: dict[str, list[dict[str, str]]] = {}
    selected_blocks_by_part: dict[str, list[dict[str, object]]] = {}
    robustness_body = _prepared_section_body(
        "Robustness And Sensitivity", sections, payload, mainline_summary
    )
    if robustness_body or evidence_rows:
        rendered_any_section = True
        document.add_heading("Robustness And Sensitivity", level=1)
        document.add_paragraph(
            "This section keeps robustness compact and decision-relevant. It uses the checks that materially affect confidence in the backtest conclusion, while avoiding a long standalone evidence dump."
        )
        if robustness_body:
            _render_narrative_section(
                document,
                "Robustness And Sensitivity",
                robustness_body,
                level=2,
                display_title="Robustness Assessment",
                include_intro=False,
            )
    if evidence_rows:
        rows_by_part: dict[str, list[dict[str, str]]] = defaultdict(list)
        for row in evidence_rows:
            rows_by_part[(row.get("part") or "").strip()].append(row)
        part_order = ["Part 1", "Part 3", "Part 4"]
        for part_label in part_order:
            part_rows = rows_by_part.get(part_label)
            if not part_rows:
                continue
            part_dir_name = _normalize_part_dir(part_label)
            part_dir = evidence_dir / part_dir_name
            enriched_rows = [dict(row, __part_dir__=str(part_dir)) for row in part_rows]
            selected_rows = _selected_rows_for_part(part_label, enriched_rows)
            appendix_rows_by_part[part_label] = selected_rows
            selected_blocks = []
            if isinstance(selected_evidence, dict):
                robustness_blocks = selected_evidence.get("robustness")
                if isinstance(robustness_blocks, list):
                    selected_blocks = [
                        block
                        for block in robustness_blocks
                        if isinstance(block, dict)
                        and str(block.get("part") or "").strip() == part_label
                    ]
            if selected_blocks:
                selected_blocks_by_part[part_label] = selected_blocks
            if not selected_blocks and not selected_rows:
                continue
            if selected_blocks:
                _render_selected_part_body(document, part_label, part_dir, selected_blocks)
            else:
                _render_part_body(document, part_label, part_dir, selected_rows)

    if robustness_body or evidence_rows:
        document.add_heading("Robustness Read-Through", level=2)
        document.add_paragraph(
            "Overall, the robustness evidence supports the main backtest conclusion across parameter-window, regime, and simulation checks. The conclusion should still be read with the fixed-window coverage caveat and the strategy's stress-period absolute-loss profile."
        )

    if not rendered_any_section:
        _add_markdown_blocks(document, _normalize_text(payload.get("analysis_text") or ""))

    if appendix_rows_by_part:
        appendix_b_started = False
        for part_label in ["Part 1", "Part 2", "Part 3", "Part 4", "Part 5"]:
            part_rows = appendix_rows_by_part.get(part_label)
            if not part_rows:
                continue
            part_dir = evidence_dir / _normalize_part_dir(part_label)
            selected_blocks = selected_blocks_by_part.get(part_label)
            if selected_blocks:
                continue
            else:
                if not _part_rows_have_appendix_assets(part_dir, part_rows):
                    continue
                if not appendix_b_started:
                    document.add_heading("Supporting Robustness Exhibits", level=2)
                    document.add_paragraph(
                        "This final subsection keeps only the extra robustness figures and tables that support the main story without repeating the main-body discussion."
                    )
                    appendix_b_started = True
                _render_part_appendix(document, part_label, part_dir, part_rows)

    output_docx.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_docx))


def main() -> None:
    parser = argparse.ArgumentParser()
    parser.add_argument("--input-json", required=True)
    parser.add_argument("--output-docx", required=True)
    args = parser.parse_args()
    payload = _load_json(Path(args.input_json))
    build_docx(payload, Path(args.output_docx))


if __name__ == "__main__":
    main()
